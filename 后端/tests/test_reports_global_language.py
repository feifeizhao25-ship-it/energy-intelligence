"""Global report exports must be natively English."""

import io
import re

import pytest

from app.routers.reports import (
    TEMPLATES,
    _demo_data,
    calc_financial_metrics,
    generate_docx,
    generate_pdf,
)


def _global_data():
    data = _demo_data()
    data.update({
        "name": "Arizona Solar Portfolio",
        "country_code": "US",
        "province_or_region": "Arizona",
        "city": "Phoenix",
        "currency": "USD",
    })
    return data


def test_global_docx_contains_no_chinese_text():
    from docx import Document
    data = _global_data()
    payload = generate_docx(
        data,
        "investment",
        TEMPLATES["investment"],
        calc_financial_metrics(data["financial"]),
        "Investment Analysis Report",
        "confidential",
        "global",
    )
    document = Document(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)
    text += "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert not re.search(r"[\u4e00-\u9fff]", text)
    assert "Assumptions and evidence" in text


def test_global_pdf_contains_no_chinese_text():
    pytest.importorskip("reportlab")
    from pypdf import PdfReader
    data = _global_data()
    payload = generate_pdf(
        data,
        "investment",
        TEMPLATES["investment"],
        calc_financial_metrics(data["financial"]),
        "Investment Analysis Report",
        "confidential",
        "global",
    )
    text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(payload)).pages)
    assert not re.search(r"[\u4e00-\u9fff]", text)
    assert "Required validation before decision" in text


@pytest.mark.parametrize("kind", ["docx", "pdf"])
def test_global_demo_export_sanitizes_untranslated_project_fields(kind):
    """A global export must remain English even when source records are CN-only."""
    data = _demo_data()
    args = (
        data,
        "investment",
        TEMPLATES["investment"],
        calc_financial_metrics(data["financial"]),
        "Investment Analysis Report",
        "confidential",
        "global",
    )
    if kind == "docx":
        from docx import Document
        document = Document(io.BytesIO(generate_docx(*args)))
        text = "\n".join(p.text for p in document.paragraphs)
    else:
        pytest.importorskip("reportlab")
        from pypdf import PdfReader
        text = "\n".join(
            page.extract_text() or ""
            for page in PdfReader(io.BytesIO(generate_pdf(*args))).pages
        )

    assert not re.search(r"[\u4e00-\u9fff]", text)
    assert "Demo Renewable Energy Project" in text
    assert re.search(
        r"City not available,\s+Region not available",
        text,
    )
