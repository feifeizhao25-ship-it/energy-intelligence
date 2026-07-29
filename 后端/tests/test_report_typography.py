"""中文版报告排版质量门禁: docx/pdf 结构断言.

覆盖:
  - docx: 封面标题 / 目录章节齐全 / 页脚页码域 (PAGE+NUMPAGES) / 表格题注
  - pdf:  页数 >= 3 / 含目录页文本 / 表格题注
英文 (global) 分支由 test_reports_global_language.py 覆盖, 本文件只测中文版.
"""

import io

import pytest

from app.routers.reports import (
    TEMPLATES,
    _demo_data,
    calc_financial_metrics,
    generate_docx,
    generate_pdf,
)

REPORT_TYPES = ["feasibility", "investment"]
COVER_TITLES = {
    "feasibility": "可行性研究报告",
    "investment": "投资分析报告",
}


def _cn_args(report_type: str):
    data = _demo_data()
    return (
        data,
        report_type,
        TEMPLATES[report_type],
        calc_financial_metrics(data["financial"]),
        "排版测试报告",
        "confidential",
        "cn",
    )


# ── DOCX ─────────────────────────────────────────────────────

@pytest.mark.parametrize("report_type", REPORT_TYPES)
def test_docx_cover_and_toc(report_type):
    """封面有报告标题, 目录页章节齐全."""
    from docx import Document

    payload = generate_docx(*_cn_args(report_type))
    document = Document(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)

    # 封面标题
    assert COVER_TITLES[report_type] in text

    # 目录页: 标题 + 全部章节条目
    assert "目  录" in text
    for i, section in enumerate(TEMPLATES[report_type].sections_zh, 1):
        assert f"第{i}章  {section}" in text, f"目录缺少章节: {section}"


@pytest.mark.parametrize("report_type", REPORT_TYPES)
def test_docx_footer_page_fields(report_type):
    """页脚统一「密级 + 第 X 页 共 Y 页」, 含 PAGE / NUMPAGES 域."""
    from docx import Document

    payload = generate_docx(*_cn_args(report_type))
    document = Document(io.BytesIO(payload))

    footer_xml = document.sections[0].footer._element.xml
    assert "PAGE" in footer_xml, "页脚缺少 PAGE 页码域"
    assert "NUMPAGES" in footer_xml, "页脚缺少 NUMPAGES 总页数域"

    footer_text = "\n".join(p.text for p in document.sections[0].footer.paragraphs)
    assert "【机密】" in footer_text, "页脚缺少中文密级标签"
    assert "第" in footer_text and "页" in footer_text and "共" in footer_text


@pytest.mark.parametrize("report_type", REPORT_TYPES)
def test_docx_table_captions_and_heading_styles(report_type):
    """表格有「表 N」编号题注; 章节标题使用统一 Heading 1 样式."""
    from docx import Document

    payload = generate_docx(*_cn_args(report_type))
    document = Document(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)

    assert "表 1" in text, "缺少表格题注 (表 1)"
    h1_paragraphs = [p for p in document.paragraphs if p.style.name == "Heading 1"]
    assert len(h1_paragraphs) >= len(TEMPLATES[report_type].sections_zh), (
        "章节标题未使用统一 Heading 1 样式"
    )


# ── PDF ──────────────────────────────────────────────────────

@pytest.mark.parametrize("report_type", REPORT_TYPES)
def test_pdf_structure(report_type):
    """PDF 页数 >= 3 (封面/目录/正文), 含目录页文本与表格题注."""
    pytest.importorskip("reportlab")
    from pypdf import PdfReader

    payload = generate_pdf(*_cn_args(report_type))
    reader = PdfReader(io.BytesIO(payload))
    assert len(reader.pages) >= 3, f"PDF 页数不足: {len(reader.pages)}"

    full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "目" in full_text and "录" in full_text, "缺少目录页文本"
    for section in TEMPLATES[report_type].sections_zh:
        assert section in full_text, f"目录/正文缺少章节: {section}"
    assert "表 1" in full_text, "缺少表格题注 (表 1)"
