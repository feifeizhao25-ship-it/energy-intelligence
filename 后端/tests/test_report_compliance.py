"""报告合规门禁: 免责声明 / 溯源表 / 双水印 / 预测标注 / 模型合规列.

覆盖运营资料硬要求:
  - 中英文报告文末固定免责声明
  - 附录「数据溯源表」 (demo 数据标「演示数据」, 缺失标「资料未及」)
  - 中文报告封面「AI 辅助生成 + 人工审核」+ 显式水印 + 隐式水印 (元数据)
  - 预测性章节显式标「这是预测，非确定」
  - Report 模型 data_sources / is_premium / reviewed / reviewed_by 落库往返
"""

import io
import re
import uuid

import pytest
from sqlalchemy import select

from app.models.database import Project
from app.models.report import Report
from app.models.user import User
from app.routers.reports import (
    DISCLAIMER_CN,
    DISCLAIMER_EN,
    GENERATOR_ID,
    TEMPLATES,
    _demo_data,
    calc_financial_metrics,
    generate_docx,
    generate_pdf,
)
from app.utils.security import get_password_hash


def _cn_args(report_type: str = "feasibility"):
    data = _demo_data()
    return (
        data,
        report_type,
        TEMPLATES[report_type],
        calc_financial_metrics(data["financial"]),
        "合规测试报告",
        "confidential",
        "cn",
    )


def _global_args(report_type: str = "investment"):
    data = _demo_data()
    data.update({
        "name": "Arizona Solar Portfolio",
        "country_code": "US",
        "province_or_region": "Arizona",
        "city": "Phoenix",
        "currency": "USD",
    })
    return (
        data,
        report_type,
        TEMPLATES[report_type],
        calc_financial_metrics(data["financial"]),
        "Investment Analysis Report",
        "confidential",
        "global",
    )


# ── 中文 DOCX ────────────────────────────────────────────────

def test_cn_docx_disclaimer_ai_label_and_provenance():
    """中文 docx: 文末固定免责声明 + 封面 AI 标注 + 附录数据溯源表."""
    from docx import Document

    payload = generate_docx(*_cn_args())
    document = Document(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )

    assert DISCLAIMER_CN in text, "缺少固定免责声明"
    assert "AI 辅助生成" in text, "封面缺少「AI 辅助生成 + 人工审核」标注"
    assert "数据溯源表" in text, "附录缺少数据溯源表章节"
    assert "演示数据" in table_text, "demo 数据未标注「演示数据」"


def test_cn_docx_forecast_explicitly_marked():
    """预测性章节显式标「这是预测，非确定」."""
    from docx import Document

    payload = generate_docx(*_cn_args())
    document = Document(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "这是预测，非确定" in text, "预测性内容缺少显式预测标注"


def test_cn_docx_dual_watermark():
    """显式水印 (页眉节级水印文字) + 隐式水印 (core properties generator+时间戳)."""
    from docx import Document

    payload = generate_docx(*_cn_args())
    document = Document(io.BytesIO(payload))

    header_xml = document.sections[0].header._element.xml
    assert "EnergyIQ" in header_xml, "缺少显式水印文字"

    core = document.core_properties
    assert core.author == GENERATOR_ID, "隐式水印缺少 generator 标识"
    assert GENERATOR_ID in (core.comments or ""), "隐式水印缺少 generator 标注"
    assert "generated_at=" in (core.comments or ""), "隐式水印缺少生成时间戳"


# ── 英文 (global) DOCX ───────────────────────────────────────

def test_global_docx_english_disclaimer_and_provenance():
    """英文 docx: 英文固定免责声明 + 英文溯源章节, 且无中文泄露."""
    from docx import Document

    payload = generate_docx(*_global_args())
    document = Document(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)
    text += "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )

    assert DISCLAIMER_EN in text, "缺少英文固定免责声明"
    assert "Data provenance" in text, "缺少英文数据溯源章节"
    assert "forecast" in text, "预测性内容缺少 forecast 标注"
    assert not re.search(r"[一-鿿]", text), "英文报告混入中文"


# ── PDF 元数据 (隐式水印) ────────────────────────────────────

def test_cn_pdf_metadata_hidden_watermark():
    """中文 PDF metadata: creator=EnergyIQ-AIGC + 生成时间戳."""
    pytest.importorskip("reportlab")
    from pypdf import PdfReader

    payload = generate_pdf(*_cn_args())
    reader = PdfReader(io.BytesIO(payload))
    meta = reader.metadata

    assert meta.creator == GENERATOR_ID, "PDF metadata 缺少 generator 标识"
    assert "generated_at=" in (meta.subject or ""), "PDF metadata 缺少生成时间戳"

    full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert DISCLAIMER_CN in full_text, "PDF 缺少固定免责声明"
    assert "数据溯源表" in full_text, "PDF 附录缺少数据溯源表"


def test_global_pdf_metadata_hidden_watermark():
    """英文 PDF metadata: creator=EnergyIQ-AIGC, 正文含英文免责."""
    pytest.importorskip("reportlab")
    from pypdf import PdfReader

    payload = generate_pdf(*_global_args())
    reader = PdfReader(io.BytesIO(payload))
    assert reader.metadata.creator == GENERATOR_ID

    full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert DISCLAIMER_EN in full_text
    assert not re.search(r"[一-鿿]", full_text), "英文 PDF 混入中文"


# ── Report 模型合规列落库往返 ────────────────────────────────

async def test_report_model_compliance_columns_roundtrip(db_session):
    """data_sources / is_premium / reviewed / reviewed_by 写入并读回."""
    user = User(
        id=str(uuid.uuid4()),
        phone=f"139{uuid.uuid4().hex[:8]}",
        password_hash=get_password_hash("TestPass123!"),
        name="终审人",
        role="user",
        market="cn",
        subscription_plan="pro",
        usage_quota={"ai_calls": {}, "report_exports": {}},
    )
    project = Project(id=str(uuid.uuid4()), user_id=user.id, name="溯源测试项目")
    report = Report(
        id=str(uuid.uuid4()),
        project_id=project.id,
        user_id=user.id,
        report_type="feasibility",
        title="合规列往返测试",
        data_sources={"capacity_mw": "演示数据", "electricity_price": "资料未及"},
        is_premium=True,
        reviewed=True,
        reviewed_by=user.id,
    )
    db_session.add_all([user, project, report])
    await db_session.commit()

    row = (
        await db_session.execute(select(Report).where(Report.id == report.id))
    ).scalar_one()
    assert row.data_sources == {"capacity_mw": "演示数据", "electricity_price": "资料未及"}
    assert row.is_premium is True
    assert row.reviewed is True
    assert row.reviewed_by == user.id


async def test_report_model_compliance_columns_defaults(db_session):
    """合规列默认值: is_premium/reviewed=False, reviewed_by/data_sources=None."""
    user = User(
        id=str(uuid.uuid4()),
        phone=f"137{uuid.uuid4().hex[:8]}",
        password_hash=get_password_hash("TestPass123!"),
        name="默认用户",
        role="user",
        market="cn",
        subscription_plan="free",
        usage_quota={"ai_calls": {}, "report_exports": {}},
    )
    project = Project(id=str(uuid.uuid4()), user_id=user.id, name="默认值测试项目")
    report = Report(
        id=str(uuid.uuid4()),
        project_id=project.id,
        user_id=user.id,
        report_type="esg",
        title="默认值测试",
    )
    db_session.add_all([user, project, report])
    await db_session.commit()

    row = (
        await db_session.execute(select(Report).where(Report.id == report.id))
    ).scalar_one()
    assert row.is_premium is False
    assert row.reviewed is False
    assert row.reviewed_by is None
    assert row.data_sources is None
