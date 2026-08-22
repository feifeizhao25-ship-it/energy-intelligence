#!/usr/bin/env python3
"""
新能源智库 - 用户场景全功能测试审计脚本 v2.0
DeepSeek-V4-Flash + MiniMax M2.7 TTS/Image + 60分制六维度评分
"""

import asyncio
import importlib.util
import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path

# ── 加载 .env（必须在 import app 之前）──────────────────────────────────────────
_env_path = Path(__file__).parent.parent / "backend" / ".env"
if _env_path.exists():
    from dotenv import load_dotenv
    load_dotenv(str(_env_path), override=True)


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "backend"))

from app.ai.llm_service import LLMService  # noqa: E402

# ── MiniMax 多模态服务（动态加载，路径含连字符）────────────────────────────────
_mm_path = PROJECT_ROOT / "backend/app-ecs/services/multimodal_service.py"
if _mm_path.exists():
    _spec = importlib.util.spec_from_file_location("multimodal_service", str(_mm_path))
    _mm = importlib.util.module_from_spec(_spec)
    _spec.loader.exec_module(_mm)
    generate_speech = getattr(_mm, "generate_speech", None)
    generate_image  = getattr(_mm, "generate_image",  None)
    generate_video  = getattr(_mm, "generate_video",  None)
else:
    generate_speech = generate_image = generate_video = None
    print("⚠️  multimodal_service.py not found")

RESULTS_PATH = Path(__file__).parent / "ai_user_scenarios_results.json"
REPORT_PATH  = PROJECT_ROOT / "docs/新能源智库-用户场景全功能测试审计报告.docx"

# ── USER_SCENARIOS v11（从 JSON 加载）─────────────────────────────────────────
_json_path = Path(__file__).parent / "user_scenarios_data.json"
with open(_json_path, encoding="utf-8") as _f:
    USER_SCENARIOS: list[dict] = json.load(_f)

# ── 60分制双语评分提示词 ────────────────────────────────────────────────────────
CURRENT_DATE = datetime.now(timezone.utc).date().isoformat()
GENERATION_MODEL = os.environ.get("SCENARIO_GENERATION_MODEL", "deepseek-v4-flash")
JUDGE_MODEL = os.environ.get("SCENARIO_JUDGE_MODEL", "qwen-plus")

JUDGE_PROMPT_CN = """你是新能源行业内容质量审计专家，请对以下AI回答进行专业评审。

【评审维度】（每项0-10分，合计60分）
1. 数据准确性（10分）：技术参数与计算数字是否准确；时效性信息是否有可核验来源和日期。
2. 计算完整性（10分）：财务/发电量等计算是否有完整推导过程，不截断。
3. 专业性（10分）：术语使用规范，标准编号与版本年份是否正确。
4. 可操作性（10分）：建议是否具体可执行，流程是否完整。
5. Skills与RAG合规性（10分）：是否显式调用[Skill:...]工具，是否引用[RAG:...]文献。
6. 时效性（10分）：以评测日期 {current_date} 为准，是否使用当前有效资料并标注发布日期/生效日期。

【输出格式】（必须严格按此格式）
总分: <X>/60
数据准确性: <X>/10 | <评价>
计算完整性: <X>/10 | <评价>
专业性: <X>/10 | <评价>
可操作性: <X>/10 | <评价>
Skills与RAG合规性: <X>/10 | <评价>
时效性: <X>/10 | <评价>
主要优点: <列出>
主要不足: <列出>
改进建议: <列出>

问题: {question}
AI回答: {answer}
"""

JUDGE_PROMPT_EN = """You are a strict renewable energy content quality auditor. Review the AI response below.

【Review Dimensions】(0-10 each, 60 total)
1. Data Accuracy (10pts): Are parameters and calculations accurate, with verifiable sources and dates for time-sensitive claims?
2. Calculation Completeness (10pts): Are financial/generation calculations fully derived?
3. Professionalism (10pts): Are terms correct? Are standard codes with version years cited?
4. Actionability (10pts): Are recommendations specific and executable?
5. Skills & RAG Compliance (10pts): Are [Skill:...] tools called? Are [RAG:...] sources cited?
6. Timeliness (10pts): As of {current_date}, does it use currently effective sources and state publication/effective dates?

【Output Format】(strictly follow)
Total Score: <X>/60
Data Accuracy: <X>/10 | <review>
Calculation Completeness: <X>/10 | <review>
Professionalism: <X>/10 | <review>
Actionability: <X>/10 | <review>
Skills & RAG Compliance: <X>/10 | <review>
Timeliness: <X>/10 | <review>
Key Strengths: <list>
Key Weaknesses: <list>
Improvement Suggestions: <list>

Question: {question}
AI Answer: {answer}
"""


# ── 核心异步函数 ────────────────────────────────────────────────────────────────
async def generate_response(service: LLMService, prompt: str, system: str, locale: str = "cn") -> str:
    messages = [{"role": "system", "content": system},
                {"role": "user",   "content": prompt}]
    chunks: list[str] = []
    target_locale = "global" if locale == "en" else "cn"
    async for chunk in service.chat_stream(messages, model=GENERATION_MODEL, target_locale=target_locale):
        if chunk.startswith("data:"):
            try:
                d = json.loads(chunk[5:].strip())
                if d.get("type") == "text":
                    chunks.append(d.get("delta", ""))
            except Exception:
                pass
    return "".join(chunks)


async def judge_response(service: LLMService, question: str, answer: str, locale: str) -> str:
    tmpl = JUDGE_PROMPT_CN if locale == "cn" else JUDGE_PROMPT_EN
    messages = [{"role": "system", "content": "You are a strict industry content quality auditor."},
                {"role": "user",   "content": tmpl.format(current_date=CURRENT_DATE, question=question, answer=answer)}]
    chunks: list[str] = []
    target_locale = "global" if locale == "en" else "cn"
    if JUDGE_MODEL == GENERATION_MODEL:
        raise RuntimeError("SCENARIO_JUDGE_MODEL must differ from SCENARIO_GENERATION_MODEL")
    async for chunk in service.chat_stream(messages, model=JUDGE_MODEL, target_locale=target_locale):
        if chunk.startswith("data:"):
            try:
                d = json.loads(chunk[5:].strip())
                if d.get("type") == "text":
                    chunks.append(d.get("delta", ""))
            except Exception:
                pass
    return "".join(chunks)


def parse_score(judge_text: str) -> int:
    for line in judge_text.split("\n"):
        if "总分" in line or "Total Score" in line:
            try:
                cleaned = line.replace("*", "").strip()
                parts = cleaned.split(":")
                if len(parts) >= 2:
                    return int(parts[1].strip().split("/")[0].strip())
            except Exception:
                pass
    return 0


# ── 多模态诊断（场景C和F触发）────────────────────────────────────────────────────
async def run_multimodal(scenario_id: str, locale: str) -> dict:
    result: dict = {"speech": None, "image": None}
    if scenario_id not in ("user-c-om-supervisor", "user-f-om-director"):
        return result

    minimax_key = os.environ.get("MINIMAX_API_KEY", "")
    if not minimax_key:
        print("    ⚠️  MINIMAX_API_KEY not set — skipping multimodal")
        return result

    if generate_speech:
        try:
            text = ("检测到第三号汇流箱第五路电流异常下降，请立即启动组串级诊断程序。"
                    if locale == "cn"
                    else "Alert: Combiner box #3 string #5 current anomaly. Initiating PV string diagnostics.")
            audio = await asyncio.to_thread(generate_speech, text,
                                            voice="male-qn-qingse", speed=1.0)
            p = PROJECT_ROOT / "uploads" / f"alert_{scenario_id}.mp3"
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_bytes(audio)
            result["speech"] = str(p)
            print(f"    🔊 Speech: {p.name}")
        except Exception as e:
            print(f"    ⚠️  Speech failed: {e}")

    if generate_image:
        try:
            img_result = await asyncio.to_thread(
                generate_image,
                "Solar PV array thermal infrared drone inspection, hotspot on string 5, professional")
            # generate_image returns dict with data.image_urls
            import urllib.request
            urls = img_result.get("data", {}).get("image_urls", [])
            if urls:
                img_url = urls[0]
                p = PROJECT_ROOT / "uploads" / f"thermal_{scenario_id}.jpeg"
                p.parent.mkdir(parents=True, exist_ok=True)
                urllib.request.urlretrieve(img_url, str(p))
                result["image"] = str(p)
                print(f"    🖼️  Image: {p.name}")
            else:
                result["image_url"] = str(img_result)
                print(f"    🖼️  Image result: {img_result}")
        except Exception as e:
            print(f"    ⚠️  Image failed: {e}")

    return result


# ── 主运行循环 ─────────────────────────────────────────────────────────────────
async def run_all() -> list[dict]:
    service = LLMService()
    results: list[dict] = []

    print("=" * 80)
    print("新能源智库 · 6-Scenario Bilingual Audit  |  DeepSeek-V4-Flash  |  60-pt Rubric")
    print(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 80)

    for idx, sc in enumerate(USER_SCENARIOS, 1):
        flag = "🇨🇳" if sc["locale"] == "cn" else "🌍"
        print(f"\n[{idx}/6] {flag} {sc['name']}  |  {sc['dimension']}")
        print("-" * 70)

        print("  ⏳ Generating answer...")
        t0 = datetime.now()
        # For EN locale, prepend language enforcement to system prompt
        system = sc["system_prompt"]
        if sc["locale"] == "en":
            system = ("CRITICAL INSTRUCTION: You MUST respond ENTIRELY in English. "
                      "Every word, sentence, and character must be in English. "
                      "Do NOT use Chinese characters. This is mandatory.\n\n") + system
        answer = await generate_response(service, sc["prompt"], system, sc["locale"])
        elapsed = (datetime.now() - t0).total_seconds()
        print(f"  ✅ {len(answer)} chars  ({elapsed:.1f}s)")

        print("  ⚖️  Judging (60-pt)...")
        judgement = await judge_response(service, sc["prompt"], answer, sc["locale"])
        score = parse_score(judgement)
        grade = "A" if score >= 54 else "B+" if score >= 48 else "B" if score >= 42 else "C"
        print(f"  📊 {score}/60  Grade: {grade}")

        mm = await run_multimodal(sc["id"], sc["locale"])

        results.append({
            "id": sc["id"], "name": sc["name"], "locale": sc["locale"],
            "dimension": sc["dimension"], "prompt": sc["prompt"],
            "answer": answer, "judgement": judgement,
            "score": score, "grade": grade,
            "multimodal": mm, "timestamp": datetime.now().isoformat(),
        })

    return results


# ── Word 报告生成 ───────────────────────────────────────────────────────────────
def _shading(cell, hex_color: str):
    from docx.oxml import parse_xml
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:fill="{hex_color}"/>'))


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def _p(doc, text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(size)
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor(*color)
    return p


def build_report(results: list[dict]):
    doc = Document()

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("新能源智库 · 用户场景全功能测试审计报告")
    r.font.size = Pt(24); r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.color.rgb = RGBColor(0x27, 0x72, 0xB0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        f"AI 引擎: DeepSeek-V4-Flash  |  多模态: MiniMax M2.7\n"
        f"评分标准: 60分制六维度  |  生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    r2.font.size = Pt(11); r2.font.name = "微软雅黑"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # Summary table
    _h(doc, "一、评分汇总", 1)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(["场景ID", "用户角色", "维度", "得分", "等级"]):
        c = tbl.rows[0].cells[i]
        c.text = h; _shading(c, "E8F4FD")
        for p in c.paragraphs:
            for rn in p.runs:
                rn.bold = True; rn.font.name = "微软雅黑"
                rn._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    gc = {"A": (0x2A, 0x9F, 0x58), "B+": (0xD8, 0x8A, 0x10),
          "B": (0xE5, 0x95, 0x00), "C": (0xC9, 0x30, 0x2C)}
    total = 0
    for r in results:
        row = tbl.add_row()
        vals = [r["id"], r["name"], r["dimension"], f"{r['score']}/60", r["grade"]]
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for rn in p.runs:
                    rn.font.name = "微软雅黑"
                    rn._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(*gc.get(r["grade"], (0,0,0)))
        total += r["score"]

    avg = total / len(results) if results else 0
    _p(doc, f"\n综合平均得分: {avg:.1f}/60  ({avg/60*100:.1f}%)", bold=True)
    _p(doc, "评分标准: A(≥54分)优秀 | B+(≥48分)良好 | B(≥42分)合格 | C(<42分)需改进")
    doc.add_page_break()

    # Individual sections
    _h(doc, "二、场景详细问答与评审", 1)
    for i, r in enumerate(results, 1):
        doc.add_page_break()
        flag = "🇨🇳" if r["locale"] == "cn" else "🌍"
        _h(doc, f"{i}. {flag} {r['name']} — {r['dimension']}", 2)
        _p(doc, f"场景ID: {r['id']}  |  语言: {'中文' if r['locale']=='cn' else 'English'}  |  "
                f"得分: {r['score']}/60  等级: {r['grade']}", bold=True,
           color=(0x27, 0x72, 0xB0))

        _p(doc, "📋 用户问题:", bold=True, color=(0x27, 0x72, 0xB0))
        _p(doc, r["prompt"], size=9.5)

        _p(doc, "🤖 AI 完整回答 (DeepSeek-V4-Flash):", bold=True, color=(0x27, 0x72, 0xB0))
        for para in r["answer"].split("\n\n"):
            if para.strip():
                _p(doc, para.strip(), size=9.5)

        _p(doc, "📊 六维度评审结果 (60分制):", bold=True, color=(0xD8, 0x8A, 0x10))
        for line in r["judgement"].split("\n"):
            if line.strip():
                _p(doc, line.strip(), size=9)

        mm = r.get("multimodal", {})
        if mm.get("speech"):
            _p(doc, f"🔊 语音诊断音频: {mm['speech']}", size=9, color=(0x27, 0x72, 0xB0))
        if mm.get("image"):
            _p(doc, "🖼️  热成像诊断图:", bold=True, color=(0x27, 0x72, 0xB0))
            try:
                doc.add_picture(mm["image"], width=Inches(4.5))
            except Exception:
                _p(doc, f"  图片路径: {mm['image']}", size=9)

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(f"\n✅ Word 报告已生成: {REPORT_PATH}")


# ── 入口 ────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    results = asyncio.run(run_all())

    with open(RESULTS_PATH, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"✅ JSON 结果: {RESULTS_PATH}")

    build_report(results)

    scores = [r["score"] for r in results]
    avg = sum(scores) / len(scores) if scores else 0
    print("\n" + "=" * 80)
    print(f"📊 最终汇总: {len(results)} 个场景  |  平均 {avg:.1f}/60 ({avg/60*100:.1f}%)")
    for r in results:
        flag = "🇨🇳" if r["locale"] == "cn" else "🌍"
        print(f"   {flag} [{r['grade']:2s}] {r['score']:2d}/60  {r['name']}")
    print("=" * 80)
