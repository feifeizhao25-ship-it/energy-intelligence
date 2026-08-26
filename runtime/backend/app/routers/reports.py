"""
报告生成路由 — 专业投资分析报告生成器
========================================
依赖: python-docx, reportlab (pip install python-docx reportlab)
可选: jinja2 (模板渲染)

功能:
  - 4 种专业报告模板: 可行性研究 / 投资分析 / 合规 / ESG
  - Word (.docx) 专业排版: 封面 + 目录 + 正文 + 表格 + 页眉页脚
  - PDF 生成: reportlab 流式布局
  - 项目数据查询 → fallback demo 数据
  - 财务指标: NPV / IRR / Payback / LCOE

颜色方案: 深蓝(#1a2744) + 金色(#c9a961)
字体: SimHei(中文标题) / SimSun(中文正文) / Arial(英文)

自测: python3 -c "from app.routers.reports import router; print(router.routes)"
"""

from __future__ import annotations

import io
import os
import math
import asyncio
import logging
import sys
from datetime import datetime, date
from enum import Enum
from typing import Any, Dict, List, Optional
from uuid import UUID, uuid4

from fastapi import (
    APIRouter,
    BackgroundTasks,
    Depends,
    HTTPException,
    Query,
    status,
)
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

# ── 项目内部依赖 ──────────────────────────────────────────────
from app.database import AsyncSessionLocal, get_db
from app.core.dependencies import get_current_user_id
from app.models.report import Report
from app.models.project import Project
from app.models.cashflow import CashflowProjection
from app.schemas.common import SuccessResponse
from app.utils.response import success
from app.services.evidence_envelope import build_envelope, source_from_dict

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/reports", tags=["报告中心"])


def _report_file(report_id: str, extension: str) -> str:
    """Return a report path only for canonical server-generated UUIDs."""
    try:
        safe_id = str(UUID(report_id))
    except ValueError as exc:
        raise HTTPException(status_code=404, detail="报告不存在或无权访问") from exc
    report_dir = os.path.realpath(os.path.join(os.getcwd(), "generated_reports"))
    return os.path.join(report_dir, f"{safe_id}.{extension}")

# ═══════════════════════════════════════════════════════════════
# 1. 常量 & 枚举
# ═══════════════════════════════════════════════════════════════

# 品牌色
COLOR_PRIMARY = "1A2744"   # 深蓝
COLOR_ACCENT = "C9A961"    # 金色
COLOR_LIGHT_GRAY = "F5F5F5"
COLOR_MEDIUM_GRAY = "888888"
COLOR_RISK_RED = "E74C3C"
COLOR_RISK_YELLOW = "F39C12"
COLOR_RISK_GREEN = "27AE60"

# 字体 fallback
try:
    from docx import Document as _DocTest
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# DOCX does not embed fonts by default. Use a font installed in the production
# image and the native macOS equivalent during render QA.
FONT_TITLE_CN = "STHeiti" if sys.platform == "darwin" else "Noto Sans CJK SC"
FONT_BODY_CN = FONT_TITLE_CN
FONT_EN = "Arial"

# ── 版式常量 (docx / pdf 两个渲染器共用, 改排版只动这里) ──
LAYOUT: Dict[str, Any] = {
    # 页面 (A4, 单位 cm)
    "page": {
        "width_cm": 21.0,
        "height_cm": 29.7,
        "margin_top_cm": 2.5,
        "margin_bottom_cm": 2.5,
        "margin_left_cm": 2.5,
        "margin_right_cm": 2.5,
    },
    # 字号阶梯 (pt)
    "font_size": {
        "cover_title": 28,
        "cover_sub": 16,
        "toc_title": 20,
        "h1": 18,
        "h2": 14,
        "h3": 12,
        "body": 11,
        "table_head": 10,
        "table_body": 9,
        "caption": 9,
        "small": 8,
    },
    # 间距阶梯 (pt)
    "spacing": {
        "body_after": 6,
        "body_line": 20,
        "h1_before": 20,
        "h1_after": 10,
        "h2_before": 14,
        "h2_after": 8,
        "caption_after": 4,
    },
}


def _confidential_label_cn(confidential: str) -> str:
    """密级英文枚举值 → 中文标签 (未知值原样返回)"""
    try:
        return ConfidentialLevel(confidential).label_cn
    except ValueError:
        return confidential


class ReportTemplateType(str, Enum):
    """报告模板类型"""
    FEASIBILITY = "feasibility"    # 可行性研究报告
    INVESTMENT = "investment"      # 投资分析报告
    COMPLIANCE = "compliance"      # 合规报告
    ESG = "esg"                    # ESG报告


class OutputFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx"


class ConfidentialLevel(str, Enum):
    PUBLIC = "public"
    INTERNAL = "internal"
    CONFIDENTIAL = "confidential"
    TOP_SECRET = "top_secret"

    # i18n labels
    @property
    def label_cn(self) -> str:
        return {"public": "公开", "internal": "内部", "confidential": "机密", "top_secret": "绝密"}[self.value]

    @property
    def label_en(self) -> str:
        return {"public": "Public", "internal": "Internal", "confidential": "Confidential", "top_secret": "Top Secret"}[self.value]


# ═══════════════════════════════════════════════════════════════
# 2. 请求 / 响应 Schema
# ═══════════════════════════════════════════════════════════════

class GenerateReportRequest(BaseModel):
    """生成报告请求"""
    project_id: str = Field(..., min_length=1, description="当前用户拥有的项目ID")
    report_type: ReportTemplateType = Field(ReportTemplateType.FEASIBILITY, description="报告模板类型")
    format: OutputFormat = Field(OutputFormat.PDF, description="输出格式 pdf/docx")
    title: Optional[str] = Field(None, description="自定义标题 (空则自动生成)")
    confidential: ConfidentialLevel = Field(ConfidentialLevel.CONFIDENTIAL, description="机密等级")


class GenerateReportResponse(BaseModel):
    """生成报告响应"""
    report_id: str
    status: str = "generating"
    report_type: str
    format: str
    download_url: Optional[str] = None
    message: str = "报告生成中，请稍后轮询状态"


class TemplateInfo(BaseModel):
    """模板信息 — 中英双语"""
    id: str
    name_zh: str
    name_en: str
    description_zh: str
    description_en: str
    sections_zh: List[str]
    sections_en: List[str]
    icon: str = "📄"

    @property
    def name(self) -> str:
        return self.name_zh  # backwards compat

    @property
    def description(self) -> str:
        return self.description_zh

    @property
    def sections(self) -> List[str]:
        return self.sections_zh


# ═══════════════════════════════════════════════════════════════
# 3. 模板定义 — 4 种报告的内容结构
# ═══════════════════════════════════════════════════════════════

TEMPLATES: Dict[str, TemplateInfo] = {
    "feasibility": TemplateInfo(
        id="feasibility",
        name_zh="可行性研究报告",
        name_en="Feasibility Study Report",
        description_zh="资源评估、技术方案、财务分析(NPV/IRR/LCOE)、风险评估、结论建议",
        description_en="Resource assessment, technical design, financial analysis (NPV/IRR/LCOE), risk assessment, recommendations",
        sections_zh=[
            "项目概述", "资源评估", "技术方案",
            "财务分析", "风险评估", "结论与建议",
        ],
        sections_en=[
            "Project Overview", "Resource Assessment", "Technical Design",
            "Financial Analysis", "Risk Assessment", "Conclusions & Recommendations",
        ],
        icon="📊",
    ),
    "investment": TemplateInfo(
        id="investment",
        name_zh="投资分析报告",
        name_en="Investment Analysis Report",
        description_zh="执行摘要、市场分析、财务模型、敏感性分析、ESG评估、投资建议",
        description_en="Executive summary, market analysis, financial modeling, sensitivity analysis, ESG assessment, investment recommendation",
        sections_zh=[
            "执行摘要", "市场分析", "财务模型",
            "敏感性分析", "ESG评估", "投资建议",
        ],
        sections_en=[
            "Executive Summary", "Market Analysis", "Financial Model",
            "Sensitivity Analysis", "ESG Assessment", "Investment Recommendation",
        ],
        icon="💰",
    ),
    "compliance": TemplateInfo(
        id="compliance",
        name_zh="合规报告",
        name_en="Compliance Report",
        description_zh="法规清单、许可证状态、环境影响评估、土地使用、电网接入",
        description_en="Regulatory checklist, permit status, environmental impact, land use, grid interconnection",
        sections_zh=[
            "法规清单", "许可证状态", "环境影响评估",
            "土地使用分析", "电网接入方案",
        ],
        sections_en=[
            "Regulatory Checklist", "Permit Status", "Environmental Impact Assessment",
            "Land Use Analysis", "Grid Interconnection Plan",
        ],
        icon="✅",
    ),
    "esg": TemplateInfo(
        id="esg",
        name_zh="ESG报告",
        name_en="ESG Report",
        description_zh="碳减排、社会影响、治理结构、可持续发展目标对接",
        description_en="Carbon reduction, social impact, governance, SDG alignment",
        sections_zh=[
            "碳排放与减排", "社会影响", "公司治理",
            "SDG 对接", "持续改进",
        ],
        sections_en=[
            "Carbon Emissions & Reduction", "Social Impact", "Corporate Governance",
            "SDG Alignment", "Continuous Improvement",
        ],
        icon="🌱",
    ),
}


# ═══════════════════════════════════════════════════════════════
# 4. 数据获取 — 项目数据 → fallback demo 数据
# ═══════════════════════════════════════════════════════════════

async def fetch_project_data(
    db: Optional[AsyncSession],
    project_id: str,
    user_id: str,
) -> Dict[str, Any]:
    """
    查询当前用户拥有的项目及财务数据。不存在、越权或查询失败均不生成报告。
    """
    if not db:
        raise HTTPException(status_code=503, detail="报告数据服务暂不可用")
    result = await db.execute(
        select(Project).where(Project.id == project_id, Project.user_id == user_id)
    )
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在或无权访问")

    cf_result = await db.execute(
        select(CashflowProjection)
        .where(CashflowProjection.project_id == project_id)
        .order_by(CashflowProjection.created_at.desc())
        .limit(1)
    )
    return _build_project_dict(project, cf_result.scalar_one_or_none())


def _build_project_dict(project: Any, cf: Any = None) -> Dict[str, Any]:
    """从 ORM 对象构建数据字典"""
    data: Dict[str, Any] = {
        "name": getattr(project, "name", "未命名项目"),
        "project_type": getattr(project, "project_type", "solar_pv"),
        "capacity_mw": float(getattr(project, "capacity_mw", 0) or 0),
        "country_code": getattr(project, "country_code", "CN"),
        "province_or_region": getattr(project, "province_or_region", ""),
        "city": getattr(project, "city", ""),
        "currency": getattr(project, "currency", "CNY"),
        "description": getattr(project, "description", ""),
        "status": getattr(project, "status", "development"),
    }

    payload = getattr(cf, "yearly_cashflows", None) if cf else None
    if isinstance(payload, dict):
        data["financial"] = dict(payload.get("assumptions") or {})
        data["financial"]["yearly_cashflows"] = payload.get("values") or []
        data["financial"]["stored_metrics"] = {
            "irr": getattr(cf, "irr", None), "npv": getattr(cf, "npv", None),
            "lcoe": getattr(cf, "lcoe", None), "payback_years": getattr(cf, "payback_years", None),
        }
        data["data_sources"] = payload.get("sources") or []
    else:
        data["financial"] = {}
        data["data_sources"] = []

    return data


REQUIRED_FINANCIAL_INPUTS = {
    "initial_investment", "project_life_years", "discount_rate",
    "annual_generation_mwh", "electricity_price", "opex_annual",
}


def validate_financial_evidence(data: Dict[str, Any]) -> None:
    """Reject investment conclusions unless assumptions and traceable sources exist."""
    financial = data.get("financial") or {}
    missing = sorted(REQUIRED_FINANCIAL_INPUTS - set(financial))
    if missing:
        raise HTTPException(status_code=422, detail=f"财务模型缺少必要输入: {', '.join(missing)}")
    if not financial.get("yearly_cashflows"):
        raise HTTPException(status_code=422, detail="财务模型缺少年度现金流序列")
    sources = data.get("data_sources") or []
    verified = [s for s in sources if isinstance(s, dict) and s.get("title") and s.get("url") and s.get("retrieved_at")]
    if not verified:
        raise HTTPException(status_code=422, detail="财务模型缺少可核验的数据来源（标题、URL、获取日期）")
    # fail-closed：已标记过期（stale）的政策类来源不得支撑投资结论，
    # 必须重新核验后才能生成报告。
    stale_policy = [
        s for s in sources
        if isinstance(s, dict) and s.get("type") == "policy" and s.get("freshness_status") == "stale"
    ]
    if stale_policy:
        titles = ", ".join(str(s.get("title", "<untitled>")) for s in stale_policy)
        raise HTTPException(
            status_code=422,
            detail=f"政策类数据来源已过期（stale），需重新核验后方可生成投资结论: {titles}",
        )


def _format_citation(index: int, src: Dict[str, Any]) -> str:
    """格式化单条数据来源引用；文献溯源字段（作者/版本/发布日期/页码）存在即透出。"""
    parts = [
        f"[{index}] {src.get('title', '未命名来源')}",
        str(src.get('url', '无URL')),
    ]
    if src.get("authors"):
        parts.append(f"作者：{src['authors']}")
    if src.get("version"):
        parts.append(f"版本：{src['version']}")
    if src.get("published_at"):
        parts.append(f"发布日期：{src['published_at']}")
    if src.get("locator"):
        parts.append(f"页码/段落：{src['locator']}")
    parts.append(f"获取日期：{src.get('retrieved_at', '未记录')}")
    if src.get("license_note"):
        parts.append(f"许可：{src['license_note']}")
    return " | ".join(parts)


def _demo_data() -> Dict[str, Any]:
    """Demo 数据用于样例报告"""
    return {
        "name": "示例光伏电站项目 (100MW)",
        "project_type": "solar_pv",
        "capacity_mw": 100.0,
        "country_code": "CN",
        "province_or_region": "甘肃省",
        "city": "酒泉市",
        "currency": "CNY",
        "description": "位于甘肃省酒泉市的100MW光伏电站项目，年等效利用小时数约1,650小时。",
        "status": "development",
        "financial": {
            "initial_investment": 360_000_000,      # 3.6亿
            "project_life_years": 25,
            "discount_rate": 0.08,
            "annual_generation_mwh": 165_000,        # 165GWh
            "electricity_price": 0.32,               # 元/kWh
            "price_escalation_rate": 0.02,
            "opex_annual": 7_200_000,                # 720万/年
            "opex_escalation_rate": 0.025,
            "loan_ratio": 0.70,
            "loan_interest_rate": 0.045,
            "loan_term_years": 15,
        },
    }


# ═══════════════════════════════════════════════════════════════
# 5. 财务指标计算
# ═══════════════════════════════════════════════════════════════

def calc_financial_metrics(fin: Dict[str, Any]) -> Dict[str, Any]:
    """
    计算 NPV / IRR / Payback / LCOE 等核心财务指标
    """
    inv = fin.get("initial_investment", 360_000_000)
    years = fin.get("project_life_years", 25)
    dr = fin.get("discount_rate", 0.08)
    gen_mwh = fin.get("annual_generation_mwh", 165_000)
    tariff = fin.get("electricity_price", 0.32)
    esc = fin.get("price_escalation_rate", 0.02)
    opex = fin.get("opex_annual", 7_200_000)
    opex_esc = fin.get("opex_escalation_rate", 0.025)

    # 年度现金流
    cashflows: List[float] = [-inv]
    cumulative = -inv
    payback_year = None

    for y in range(1, years + 1):
        revenue = gen_mwh * 1000 * tariff * ((1 + esc) ** (y - 1))
        cost = opex * ((1 + opex_esc) ** (y - 1))
        net = revenue - cost
        cashflows.append(net)

        cumulative += net
        if payback_year is None and cumulative >= 0:
            payback_year = y

    # NPV
    npv = sum(cf / ((1 + dr) ** i) for i, cf in enumerate(cashflows))

    # IRR (Newton-Raphson)
    irr = _calc_irr(cashflows)

    # LCOE (元/kWh)
    total_gen_kwh = sum(gen_mwh * 1000 * ((1 + 0.005) ** (y - 1)) for y in range(1, years + 1))
    total_cost_pv = inv + sum(
        opex * ((1 + opex_esc) ** (y - 1)) / ((1 + dr) ** y) for y in range(1, years + 1)
    )
    lcoe = total_cost_pv / total_gen_kwh if total_gen_kwh > 0 else 0

    # ROI
    total_net = sum(cashflows[1:])
    roi = total_net / inv if inv > 0 else 0

    return {
        "npv": npv,
        "irr": irr,
        "payback_years": payback_year or years,
        "lcoe": lcoe,
        "roi": roi,
        "initial_investment": inv,
        "annual_revenue_year1": gen_mwh * 1000 * tariff,
        "annual_opex_year1": opex,
        "project_life": years,
        "cashflows": cashflows,
    }


def _calc_irr(cashflows: List[float], max_iter: int = 100, tol: float = 1e-6) -> float:
    """Newton-Raphson 法求 IRR"""
    rate = 0.10  # 初始猜测 10%
    for _ in range(max_iter):
        npv = sum(cf / ((1 + rate) ** i) for i, cf in enumerate(cashflows))
        dnpv = sum(
            -i * cf / ((1 + rate) ** (i + 1)) for i, cf in enumerate(cashflows)
        )
        if abs(dnpv) < 1e-12:
            break
        rate_new = rate - npv / dnpv
        if abs(rate_new - rate) < tol:
            return rate_new
        rate = rate_new
        if rate < -0.99:
            rate = -0.99
    return rate


def fmt_money(v: float, unit: str = "元", decimals: int = 0) -> str:
    """千分位金额格式化"""
    if abs(v) >= 100_000_000:
        return f"{v / 100_000_000:,.2f} 亿{unit}"
    elif abs(v) >= 10_000:
        return f"{v / 10_000:,.2f} 万{unit}"
    return f"{v:,.{decimals}f} {unit}"


def fmt_pct(v: float, decimals: int = 1) -> str:
    """百分比格式化 — 全报告统一 1 位小数"""
    return f"{v * 100:.{decimals}f}%"


# ═══════════════════════════════════════════════════════════════
# 6. Word 生成器 (python-docx)
# ═══════════════════════════════════════════════════════════════

def _global_section_body(section: str, data: Dict[str, Any], metrics: Dict[str, Any]) -> str:
    """English body copy for one report section (global market)."""
    name = data.get("name", "the project")
    capacity = data.get("capacity_mw", "-")
    city = data.get("city", "")
    region = data.get("province_or_region", "")
    irr = metrics.get("irr")
    npv = metrics.get("npv")
    lcoe = metrics.get("lcoe")
    irr_txt = f"{irr:.2%}" if isinstance(irr, (int, float)) else "n/a"
    npv_txt = f"{npv:,.0f}" if isinstance(npv, (int, float)) else "n/a"
    lcoe_txt = f"{lcoe:.4f}" if isinstance(lcoe, (int, float)) else "n/a"
    bodies = {
        "Executive Summary": (
            f"{name} is a {capacity} MW renewable energy project located in "
            f"{city}, {region}. Key indicators: IRR {irr_txt}, NPV {npv_txt} "
            f"{data.get('currency', 'USD')}, LCOE {lcoe_txt}."
        ),
        "Market Analysis": (
            "Market conditions are summarized from public sources cited in the "
            "assumptions register; figures are estimates, not guarantees."
        ),
        "Financial Model": (
            f"The model spans {data.get('financial', {}).get('project_life_years', 25)} years "
            f"with a discount rate of {data.get('financial', {}).get('discount_rate', 0.08):.1%}."
        ),
        "Sensitivity Analysis": (
            "Sensitivity is assessed on electricity price, capex and capacity "
            "factor within +/-20% bands."
        ),
        "ESG Assessment": (
            "ESG screening covers land use, grid impact and community "
            "consultation status as declared in the project register."
        ),
        "Investment Recommendation": (
            "The recommendation is conditional on the validation items listed "
            "below being completed by qualified parties."
        ),
    }
    return bodies.get(section, f"{section}: see project documentation for {name}.")


def _generate_docx_global(
    data: Dict[str, Any],
    report_type: str,
    template_info: TemplateInfo,
    metrics: Dict[str, Any],
    title: str,
    confidential: str,
) -> bytes:
    """Global-market DOCX: fully English output (no untranslated Chinese)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Confidentiality: {confidential}")
    doc.add_paragraph(
        f"Project: {data.get('name', '-')} | Capacity: {data.get('capacity_mw', '-')} MW | "
        f"Location: {data.get('city', '-')}, {data.get('province_or_region', '-')}, "
        f"{data.get('country_code', '-')} | Currency: {data.get('currency', 'USD')}"
    )

    for section in template_info.sections_en:
        doc.add_heading(section, level=1)
        doc.add_paragraph(_global_section_body(section, data, metrics))

    doc.add_heading("Assumptions and evidence", level=1)
    doc.add_paragraph(
        "All figures in this report are estimates derived from the assumptions "
        "register. Each key input lists its evidence source and verification date."
    )

    doc.add_heading("Required validation before decision", level=1)
    doc.add_paragraph(
        "Independent engineering review, grid interconnection confirmation and "
        "legal due diligence must be completed before any investment decision."
    )

    doc.add_paragraph(
        "Disclaimer: forecasts depend on assumptions; actual results may differ. "
        "The authors accept no liability for decisions made on this report."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _generate_pdf_global(
    data: Dict[str, Any],
    report_type: str,
    template_info: TemplateInfo,
    metrics: Dict[str, Any],
    title: str,
    confidential: str,
) -> bytes:
    """Global-market PDF: fully English output (no untranslated Chinese)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from reportlab.lib.units import mm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=title,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph(title, styles["Title"]),
        Spacer(1, 6 * mm),
        Paragraph(f"Confidentiality: {confidential}", styles["Normal"]),
        Paragraph(
            f"Project: {data.get('name', '-')} | Capacity: {data.get('capacity_mw', '-')} MW | "
            f"Location: {data.get('city', '-')}, {data.get('province_or_region', '-')}, "
            f"{data.get('country_code', '-')} | Currency: {data.get('currency', 'USD')}",
            styles["Normal"],
        ),
        Spacer(1, 6 * mm),
    ]
    for section in template_info.sections_en:
        story.append(Paragraph(section, styles["Heading1"]))
        story.append(Paragraph(_global_section_body(section, data, metrics), styles["Normal"]))
        story.append(Spacer(1, 3 * mm))

    story.append(Paragraph("Assumptions and evidence", styles["Heading1"]))
    story.append(Paragraph(
        "All figures in this report are estimates derived from the assumptions "
        "register. Each key input lists its evidence source and verification date.",
        styles["Normal"],
    ))
    story.append(Paragraph("Required validation before decision", styles["Heading1"]))
    story.append(Paragraph(
        "Independent engineering review, grid interconnection confirmation and "
        "legal due diligence must be completed before any investment decision.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph(
        "Disclaimer: forecasts depend on assumptions; actual results may differ. "
        "The authors accept no liability for decisions made on this report.",
        styles["Italic"],
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def generate_docx(
    data: Dict[str, Any],
    report_type: str,
    template_info: TemplateInfo,
    metrics: Dict[str, Any],
    title: str,
    confidential: str,
    market: str = "cn",
) -> bytes:
    """
    使用 python-docx 生成专业 Word 文档
    """
    if market in ("global", "en", "int"):
        return _generate_docx_global(data, report_type, template_info, metrics, title, confidential)
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # ── 页面设置 ──
    page = LAYOUT["page"]
    section = doc.sections[0]
    section.page_width = Cm(page["width_cm"])
    section.page_height = Cm(page["height_cm"])
    section.top_margin = Cm(page["margin_top_cm"])
    section.bottom_margin = Cm(page["margin_bottom_cm"])
    section.left_margin = Cm(page["margin_left_cm"])
    section.right_margin = Cm(page["margin_right_cm"])

    # ── 页眉页脚 ──
    _setup_header_footer(doc, title, confidential)

    # ── 设置默认字体 ──
    fs = LAYOUT["font_size"]
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_EN
    font.size = Pt(fs["body"])
    _set_run_fonts_style(style, FONT_BODY_CN)

    # ── 标题 1/2/3 样式统一 (字号/字重/颜色, 供 TOC 域识别) ──
    _setup_heading_styles(doc)

    # ════ 封面 ════
    _add_cover_page(doc, title, report_type, confidential, data)

    # ════ 目录 ════
    _add_toc_page(doc, template_info)

    # ════ 正文 ════（目录函数已写入分页符）

    if report_type == "feasibility":
        _add_feasibility_content(doc, data, metrics)
    elif report_type == "investment":
        _add_investment_content(doc, data, metrics)
    elif report_type == "compliance":
        _add_compliance_content(doc, data, metrics)
    elif report_type == "esg":
        _add_esg_content(doc, data, metrics)

    # ════ 附录 ════
    doc.add_page_break()
    _add_appendix(doc, data, metrics)

    # ── 导出 ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _setup_heading_styles(doc):
    """统一 Heading 1/2/3 样式的字号、字重、颜色与中文字体。

    正文标题通过段落样式引用 (而非逐段硬编码)，
    同时让 Word 原生 TOC 域 (\\o "1-3") 能识别章节层级。
    """
    from docx.shared import Pt, RGBColor

    fs = LAYOUT["font_size"]
    sp = LAYOUT["spacing"]
    specs = {
        "Heading 1": (fs["h1"], COLOR_PRIMARY, sp["h1_before"], sp["h1_after"]),
        "Heading 2": (fs["h2"], "2C3E50", sp["h2_before"], sp["h2_after"]),
        "Heading 3": (fs["h3"], "344A5E", sp["h2_before"] // 2, sp["h2_after"] // 2),
    }
    for style_name, (size, color, before, after) in specs.items():
        st = doc.styles[style_name]
        st.font.name = FONT_EN
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        _set_run_fonts_style(st, FONT_TITLE_CN)


def _add_field_run(paragraph, instruction: str):
    """向段落追加一个 Word 域 (如 PAGE / NUMPAGES)"""
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    run = paragraph.add_run()
    run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run2 = paragraph.add_run()
    run2._element.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> {instruction} </w:instrText>'
    ))
    run3 = paragraph.add_run()
    run3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def _setup_header_footer(doc, title: str, confidential: str):
    """页眉页脚 — 页眉报告名; 页脚统一「密级 + 第 X 页 共 Y 页」(页码域)"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]

    # 页眉 — 报告名
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = title
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        run.font.size = Pt(LAYOUT["font_size"]["small"])
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = FONT_EN

    # 页脚 — 「【密级】  第 X 页 共 Y 页」(PAGE / NUMPAGES 域, 打开 Word 自动更新)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run(f"【{_confidential_label_cn(confidential)}】  第 ")
    _add_field_run(fp, "PAGE")
    fp.add_run(" 页 共 ")
    _add_field_run(fp, "NUMPAGES")
    fp.add_run(" 页")

    for r in fp.runs:
        r.font.size = Pt(LAYOUT["font_size"]["small"])
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _set_run_fonts(run, cn_font: str = FONT_BODY_CN):
    """Set every OOXML font slot so headless renderers cannot select a missing font."""
    from docx.oxml.ns import qn
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for slot in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(slot), cn_font)
    run.font.name = cn_font


def _set_run_fonts_style(style, cn_font: str = FONT_BODY_CN):
    """Set every OOXML font slot on a style."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for slot in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(slot), cn_font)
    style.font.name = cn_font


def _add_cover_page(doc, title: str, report_type: str, confidential: str, data: Dict):
    """封面页"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    fs = LAYOUT["font_size"]

    # 空行留白
    for _ in range(4):
        doc.add_paragraph()

    # 顶部装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)

    # 报告类型
    type_names = {
        "feasibility": "可行性研究报告",
        "investment": "投资分析报告",
        "compliance": "合规审查报告",
        "esg": "ESG 专项报告",
    }
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(type_names.get(report_type, "项目报告"))
    run.font.size = Pt(fs["cover_title"])
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)
    run.font.name = FONT_EN
    _set_run_fonts(run, FONT_TITLE_CN)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get("name", "新能源项目"))
    run.font.size = Pt(fs["cover_sub"])
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_fonts(run, FONT_BODY_CN)

    # 空行
    for _ in range(6):
        doc.add_paragraph()

    # 信息块
    info_items = [
        ("项目名称", data.get("name", "—")),
        ("项目类型", _project_type_label(data.get("project_type", ""))),
        ("装机容量", f"{data.get('capacity_mw', 0):.0f} MW"),
        ("项目地点", _format_location(data)),
        ("报告日期", datetime.now().strftime("%Y年%m月%d日")),
        ("机密等级", f"【{_confidential_label_cn(confidential)}】"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(fs["body"])
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _set_run_fonts(run, FONT_BODY_CN)

    # 底部装饰线
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)

    doc.add_page_break()


def _add_toc_page(doc, template_info: TemplateInfo):
    """目录页 — Word 原生 TOC 域 (可右键更新页码)"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 标题
    p = doc.add_paragraph()
    run = p.add_run("目  录" if hasattr(template_info, 'name_zh') else "Table of Contents")
    run.font.size = Pt(LAYOUT["font_size"]["toc_title"])
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)
    _set_run_fonts(run, FONT_TITLE_CN)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Word 原生 TOC 域代码 (打开文档后右键→更新域→更新整个目录)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "右键此处 → 更新域 可生成带页码的目录"
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    run._element.append(fldChar4)

    # 静态章节预览 (非页码版)
    doc.add_paragraph()
    for i, section_name in enumerate(template_info.sections, 1):
        p = doc.add_paragraph()
        # 添加书签便于导航
        run = p.add_run(f"第{i}章  {section_name}")
        run.font.size = Pt(12)
        _set_run_fonts(run, FONT_BODY_CN)

    # 附录
    p = doc.add_paragraph()
    run = p.add_run("附录")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_run_fonts(run, FONT_BODY_CN)

    doc.add_page_break()


def _add_watermark(doc, text: str = "CONFIDENTIAL"):
    """添加水印到所有页面 (45° 旋转, 10% 透明度)"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        # 在 header 中添加水印 (使用 VML)
        watermark_xml = f'''
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:r>
                <w:pict>
                    <v:shape xmlns:v="urn:schemas-microsoft-com:vml" 
                             id="PowerPlusWaterMarkObject" 
                             o:spid="#_x0000_s2049" 
                             type="#_x0000_t136" 
                             style="position:absolute;margin-left:0;margin-top:0;width:468pt;height:117pt;rotation:315;z-index:-251654144;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin"
                             o:allowincell="f"
                             fillcolor="#C0C0C0" 
                             stroked="f">
                        <v:fill opacity=".1"/>
                        <v:textpath style="font-family:&quot;Arial&quot;;font-size:1pt" text="{text}"/>
                    </v:shape>
                </w:pict>
            </w:r>
        </w:p>'''
        # 解析 XML 并添加到 header
        from lxml import etree
        header_para = header.add_paragraph()
        header_para_element = header_para._element
        parsed = etree.fromstring(watermark_xml)
        header_para_element.append(parsed)


def _add_heading(doc, text: str, level: int = 1):
    """添加标题 — 使用统一的 Heading 1/2/3 段落样式 (字号字重见 _setup_heading_styles)"""
    from docx.shared import Pt, RGBColor

    level = max(1, min(level, 3))
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    _set_run_fonts(run, FONT_TITLE_CN)
    if level == 1:
        # 下方金色装饰线
        p2 = doc.add_paragraph()
        r2 = p2.add_run("━" * 20)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)
    return p


def _add_body(doc, text: str, bold: bool = False):
    """正文段落"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(LAYOUT["font_size"]["body"])
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_run_fonts(run, FONT_BODY_CN)
    p.paragraph_format.space_after = Pt(LAYOUT["spacing"]["body_after"])
    p.paragraph_format.line_spacing = Pt(LAYOUT["spacing"]["body_line"])
    return p


def _next_caption_number(doc, kind: str = "表") -> int:
    """图表题注自动编号 (计数器挂在 doc 实例上, 每份文档独立)"""
    counters = getattr(doc, "_caption_counters", None)
    if counters is None:
        counters = {"图": 0, "表": 0}
        doc._caption_counters = counters
    counters[kind] = counters.get(kind, 0) + 1
    return counters[kind]


def _add_caption(doc, kind: str, text: str):
    """图表题注段落: 「表 1  xxx」/「图 1  xxx」, 居中灰小字"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    n = _next_caption_number(doc, kind)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{kind} {n}  {text}")
    run.font.size = Pt(LAYOUT["font_size"]["caption"])
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _set_run_fonts(run, FONT_BODY_CN)
    p.paragraph_format.space_after = Pt(LAYOUT["spacing"]["caption_after"])
    return p


def _add_styled_table(
    doc,
    headers: List[str],
    rows: List[List[str]],
    col_widths: Optional[List[float]] = None,
    caption: Optional[str] = None,
):
    """专业表格 (深蓝表头 + 斑马纹), caption 提供时自动加「表 N」题注"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    fs = LAYOUT["font_size"]
    if caption:
        _add_caption(doc, "表", caption)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(fs["table_head"])
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_run_fonts(run, FONT_TITLE_CN)
        p.alignment = 1  # CENTER
        # 背景色
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{COLOR_PRIMARY}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(fs["table_body"])
            _set_run_fonts(run, FONT_BODY_CN)
            # 斑马纹
            if row_idx % 2 == 1:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{COLOR_LIGHT_GRAY}" w:val="clear"/>'
                )
                cell._element.get_or_add_tcPr().append(shading)

    # 表格边框
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # 表后空行
    return table


def _add_risk_matrix(doc, risks: List[Dict[str, str]]):
    """风险矩阵 (红/黄/绿)"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    _add_heading(doc, "风险矩阵", level=2)

    color_map = {
        "高": COLOR_RISK_RED,
        "中": COLOR_RISK_YELLOW,
        "低": COLOR_RISK_GREEN,
    }

    for risk in risks:
        level = risk.get("level", "中")
        color = color_map.get(level, COLOR_RISK_YELLOW)

        p = doc.add_paragraph()
        run = p.add_run(f"  ● {risk['name']}  ")
        run.font.size = Pt(11)
        run.font.bold = True
        _set_run_fonts(run, FONT_BODY_CN)

        run2 = p.add_run(f"[{level}]")
        run2.font.size = Pt(10)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(
            int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        )

        run3 = p.add_run(f" — {risk.get('description', '')}")
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _set_run_fonts(run3, FONT_BODY_CN)


# ── 可行性研究报告 正文 ──

def _add_feasibility_content(doc, data: Dict, metrics: Dict):
    """可行性研究报告"""
    fin = data.get("financial", {})
    cap = data.get("capacity_mw", 0)

    # 第1章 项目概述
    _add_heading(doc, "第一章  项目概述", level=1)

    _add_heading(doc, "1.1 项目背景", level=2)
    _add_body(doc,
        f"本报告对「{data['name']}」进行全面的可行性研究分析。"
        f"项目位于{_format_location(data)}，规划装机容量 {cap:.0f} MW，"
        f"属于{_project_type_label(data['project_type'])}类别。"
        f"项目当前状态为「{_project_status_label(data.get('status', ''))}」。"
    )

    _add_heading(doc, "1.2 项目目标", level=2)
    _add_body(doc, data.get("description", "建设新能源发电项目，实现清洁能源并网发电。"))

    _add_heading(doc, "1.3 项目基本信息", level=2)
    _add_styled_table(
        doc,
        ["项目参数", "数值"],
        [
            ["项目名称", data["name"]],
            ["项目类型", _project_type_label(data["project_type"])],
            ["装机容量", f"{cap:.0f} MW"],
            ["项目地点", _format_location(data)],
            ["预计寿命", f"{fin.get('project_life_years', 25)} 年"],
            ["货币单位", data.get("currency", "CNY")],
        ],
        col_widths=[6, 10],
        caption="项目基本信息一览",
    )

    # 第2章 资源评估
    _add_heading(doc, "第二章  资源评估", level=1)

    _add_heading(doc, "2.1 太阳能/风能资源概况", level=2)
    annual_gen = fin.get("annual_generation_mwh", 0)
    utilization_hours = annual_gen / cap if cap > 0 else 0
    _add_body(doc,
        f"根据气象数据评估，项目场址年等效利用小时数约 {utilization_hours:,.0f} 小时，"
        f"首年预计发电量 {annual_gen:,.0f} MWh。"
        f"资源等级评定为「良好」，具备开发价值。"
    )

    _add_heading(doc, "2.2 发电量预测", level=2)
    gen_rows = []
    for y in [1, 5, 10, 15, 20, 25]:
        if y <= fin.get("project_life_years", 25):
            yearly_gen = annual_gen * ((1 + 0.005) ** (y - 1))
            gen_rows.append([f"第 {y} 年", f"{yearly_gen:,.0f} MWh"])
    _add_styled_table(
        doc, ["年份", "预计发电量"], gen_rows, col_widths=[6, 10],
        caption="分年度发电量预测",
    )

    # 第3章 技术方案
    _add_heading(doc, "第三章  技术方案", level=1)

    _add_heading(doc, "3.1 系统设计", level=2)
    _add_body(doc,
        f"推荐采用分块发电、集中并网方案。"
        f"主设备包括光伏组件/风力发电机组、逆变器、升压变压器及监控系统。"
        f"总装机容量 {cap:.0f} MW，建议配置适当比例的储能系统以提高消纳率。"
    )

    _add_heading(doc, "3.2 主要设备选型", level=2)
    _add_styled_table(
        doc, ["设备类别", "规格建议", "数量"],
        [
            ["发电组件", "高效单晶硅 / 双面组件", f"{cap:.0f} MW"],
            ["逆变器", "组串式 / 集中式逆变器", f"{cap * 1.1:.0f} MW"],
            ["升压站", "35kV/110kV 升压站", "1 座"],
            ["监控系统", "SCADA 远程监控", "1 套"],
        ],
        col_widths=[5, 8, 4],
        caption="主要设备选型建议",
    )

    # 第4章 财务分析
    _add_heading(doc, "第四章  财务分析", level=1)

    _add_heading(doc, "4.1 投资估算", level=2)
    _add_styled_table(
        doc, ["财务指标", "数值"],
        [
            ["初始总投资", fmt_money(metrics["initial_investment"])],
            ["单位千瓦投资", fmt_money(metrics["initial_investment"] / cap if cap > 0 else 0) + "/kW"],
            ["首年运维费用", fmt_money(metrics["annual_opex_year1"])],
            ["首年发电收入", fmt_money(metrics["annual_revenue_year1"])],
        ],
        col_widths=[6, 10],
        caption="投资估算汇总",
    )

    _add_heading(doc, "4.2 核心财务指标", level=2)
    irr_display = fmt_pct(metrics["irr"])
    _add_styled_table(
        doc, ["指标", "数值", "评价"],
        [
            ["净现值 (NPV)", fmt_money(metrics["npv"]),
             "✓ 达标" if metrics["npv"] > 0 else "✗ 不达标"],
            ["内部收益率 (IRR)", irr_display,
             "✓ 良好" if metrics["irr"] > 0.08 else "⚠ 一般"],
            ["投资回收期", f"{metrics['payback_years']:.1f} 年",
             "✓ 合理" if metrics["payback_years"] < 10 else "⚠ 偏长"],
            ["LCOE 度电成本", f"{metrics['lcoe']:.4f} 元/kWh",
             "✗ 偏高" if metrics["lcoe"] > 0.35 else "✓ 合理"],
            ["投资回报率 (ROI)", fmt_pct(metrics["roi"]), "—"],
        ],
        col_widths=[5, 5, 6],
        caption="核心财务指标评价",
    )

    _add_heading(doc, "4.3 现金流预测 (前10年)", level=2)
    cf_rows = []
    for y in range(1, min(11, len(metrics["cashflows"]))):
        cf = metrics["cashflows"][y]
        cf_rows.append([f"第 {y} 年", fmt_money(cf)])
    _add_styled_table(doc, ["年度", "净现金流"], cf_rows, col_widths=[6, 10],
                      caption="前10年净现金流预测")

    # 第5章 风险评估
    _add_heading(doc, "第五章  风险评估", level=1)
    _add_risk_matrix(doc, [
        {"name": "政策风险", "level": "中", "description": "补贴政策调整、电价市场化影响收入"},
        {"name": "资源风险", "level": "低", "description": "气象条件波动导致发电量偏差"},
        {"name": "技术风险", "level": "低", "description": "设备故障率、衰减率超出预期"},
        {"name": "资金风险", "level": "中", "description": "利率上升导致融资成本增加"},
        {"name": "市场风险", "level": "中", "description": "电力市场化交易电价不确定性"},
    ])

    # 第6章 结论
    _add_heading(doc, "第六章  结论与建议", level=1)
    npv_ok = metrics["npv"] > 0
    irr_ok = metrics["irr"] > 0.08
    _add_body(doc,
        f"综合以上分析，项目净现值为 {fmt_money(metrics['npv'])}，"
        f"内部收益率为 {irr_display}，投资回收期为 {metrics['payback_years']:.1f} 年。"
    )
    if npv_ok and irr_ok:
        _add_body(doc,
            "结论：项目各项财务指标均达到投资要求，财务可行性良好，"
            "建议推进项目开发。",
            bold=True,
        )
    else:
        _add_body(doc,
            "结论：项目部分财务指标未达预期，建议优化方案后重新评估。",
            bold=True,
        )


# ── 投资分析报告 正文 ──

def _add_investment_content(doc, data: Dict, metrics: Dict):
    """投资分析报告"""
    cap = data.get("capacity_mw", 0)

    _add_heading(doc, "第一章  执行摘要", level=1)
    _add_body(doc,
        f"本报告对「{data['name']}」进行投资价值分析。"
        f"项目装机容量 {cap:.0f} MW，总投资 {fmt_money(metrics['initial_investment'])}。"
        f"经测算，项目 IRR 为 {fmt_pct(metrics['irr'])}，"
        f"NPV 为 {fmt_money(metrics['npv'])}，"
        f"投资回收期 {metrics['payback_years']:.1f} 年。"
    )

    _add_heading(doc, "第二章  市场分析", level=1)
    _add_heading(doc, "2.1 行业概况", level=2)
    _add_body(doc,
        "在国家「双碳」目标驱动下，新能源行业持续高速增长。"
        "2024年全国光伏新增装机超 200GW，风电新增超 70GW。"
        "预计 2025-2030 年复合增长率维持在 15-20%。"
    )
    _add_heading(doc, "2.2 电价趋势", level=2)
    _add_body(doc,
        "随着电力市场化改革推进，新能源上网电价逐步从标杆电价向市场化电价过渡。"
        "绿电交易和绿证制度为新能源项目提供额外溢价空间。"
    )

    _add_heading(doc, "第三章  财务模型", level=1)
    _add_styled_table(
        doc, ["核心指标", "数值", "行业基准"],
        [
            ["NPV", fmt_money(metrics["npv"]), "> 0"],
            ["IRR", fmt_pct(metrics["irr"]), "≥ 8%"],
            ["回收期", f"{metrics['payback_years']:.1f} 年", "≤ 10 年"],
            ["LCOE", f"{metrics['lcoe']:.4f} 元/kWh", "≤ 0.35 元/kWh"],
            ["ROI", fmt_pct(metrics["roi"]), "≥ 50%"],
        ],
        col_widths=[5, 5, 6],
        caption="财务模型核心指标与行业基准",
    )

    _add_heading(doc, "第四章  敏感性分析", level=1)
    _add_body(doc, "以下分析关键变量变动 ±10% 对 IRR 的影响：")
    base_irr = metrics["irr"]
    fin = data.get("financial", {})
    tariff = fin.get("electricity_price", 0.32)

    sens_rows = []
    for delta in [-0.10, -0.05, 0, 0.05, 0.10]:
        adjusted_tariff = tariff * (1 + delta)
        adjusted_fin = {**fin, "electricity_price": adjusted_tariff}
        adjusted_metrics = calc_financial_metrics(adjusted_fin)
        label = f"{delta:+.1%}"
        sens_rows.append([label, fmt_pct(adjusted_metrics["irr"]),
                         f"{(adjusted_metrics['irr'] - base_irr):+.1%}"])
    _add_styled_table(
        doc, ["电价变动", "调整后 IRR", "IRR 变化"],
        sens_rows, col_widths=[5, 5, 6],
        caption="电价敏感性分析 (±10%)",
    )

    _add_heading(doc, "第五章  ESG 评估", level=1)
    annual_gen_kwh = fin.get("annual_generation_mwh", 0) * 1000
    co2_reduction = annual_gen_kwh * 0.585  # kg CO2/kWh
    _add_body(doc,
        f"项目年发电量约 {annual_gen_kwh / 1e6:,.1f} GWh，"
        f"年等效减排 CO₂ 约 {co2_reduction / 1e6:,.2f} 万吨。"
        "符合国家碳中和战略方向，ESG 评级为 A 级。"
    )

    _add_heading(doc, "第六章  投资建议", level=1)
    _add_body(doc,
        f"综合财务模型和风险分析，项目 IRR ({fmt_pct(metrics['irr'])}) "
        f"{'高于' if metrics['irr'] > 0.08 else '低于'}行业基准 (8%)，"
        f"NPV {'为正' if metrics['npv'] > 0 else '为负'}。"
    )
    recommendation = (
        "建议：推荐投资。项目财务指标良好，政策风险可控，ESG 效益显著。"
        if metrics["npv"] > 0 and metrics["irr"] > 0.08
        else "建议：暂缓投资。需优化成本结构或提升电价水平后再行评估。"
    )
    _add_body(doc, recommendation, bold=True)


# ── 合规报告 正文 ──

def _add_compliance_content(doc, data: Dict, metrics: Dict):
    """合规报告"""

    _add_heading(doc, "第一章  法规清单", level=1)
    _add_styled_table(
        doc, ["法规名称", "适用范围", "合规状态"],
        [
            ["《可再生能源法》", "全国", "✓ 符合"],
            ["《电力法》", "全国", "✓ 符合"],
            ["《环境保护法》", "全国", "✓ 符合"],
            ["《土地管理法》", "全国", "⚠ 待确认"],
            ["地方新能源规划", data.get("province_or_region", ""), "✓ 符合"],
        ],
        col_widths=[6, 4, 5],
        caption="适用法规合规性清单",
    )

    _add_heading(doc, "第二章  许可证状态", level=1)
    _add_styled_table(
        doc, ["许可证/批文", "主管部门", "状态", "预计取得时间"],
        [
            ["项目核准/备案", "发改委", "待办理", "T+3 月"],
            ["环评批复", "生态环境局", "待办理", "T+6 月"],
            ["土地预审", "自然资源局", "待办理", "T+4 月"],
            ["电网接入意见", "电网公司", "待办理", "T+5 月"],
            ["施工许可证", "住建局", "待办理", "T+8 月"],
        ],
        col_widths=[5, 4, 3, 4],
        caption="许可证/批文办理状态",
    )

    _add_heading(doc, "第三章  环境影响评估", level=1)
    _add_body(doc,
        "项目建设和运营对环境的影响主要为：施工期噪声和扬尘、"
        "运营期电磁辐射和生态影响。"
        "建议委托具有资质的环评机构开展详细评估，并编制水土保持方案。"
    )

    _add_heading(doc, "第四章  土地使用分析", level=1)
    _add_body(doc,
        f"项目选址位于{_format_location(data)}，"
        "需确认土地性质（未利用地/农用地/建设用地），"
        "办理建设用地审批手续。建议优先使用未利用地或戈壁荒漠。"
    )

    _add_heading(doc, "第五章  电网接入方案", level=1)
    cap = data.get("capacity_mw", 0)
    _add_body(doc,
        f"项目装机 {cap:.0f} MW，建议接入 110kV 或 220kV 变电站。"
        "需与当地电网公司签订并网协议，明确接入系统方案、"
        "计量点及电价结算方式。"
    )


# ── ESG 报告 正文 ──

def _add_esg_content(doc, data: Dict, metrics: Dict):
    """ESG 报告"""
    fin = data.get("financial", {})
    cap = data.get("capacity_mw", 0)
    annual_gen_kwh = fin.get("annual_generation_mwh", 0) * 1000

    _add_heading(doc, "第一章  碳排放与减排", level=1)
    co2_reduction_tons = annual_gen_kwh * 0.585 / 1000  # tons
    coal_equivalent = annual_gen_kwh * 0.3 / 1000  # kg → tons standard coal

    _add_styled_table(
        doc, ["减排指标", "年度值", "25年累计"],
        [
            ["发电量", f"{annual_gen_kwh / 1e6:,.1f} GWh",
             f"{annual_gen_kwh / 1e6 * 25 * 0.98:,.0f} GWh"],
            ["CO₂ 减排", f"{co2_reduction_tons / 1e4:,.2f} 万吨",
             f"{co2_reduction_tons * 25 / 1e4 * 0.98:,.0f} 万吨"],
            ["标煤替代", f"{coal_equivalent / 1e4:,.2f} 万吨",
             f"{coal_equivalent * 25 / 1e4 * 0.98:,.0f} 万吨"],
            ["SO₂ 减排", f"{annual_gen_kwh * 0.00003 / 1e4:,.2f} 万吨", "—"],
            ["NOx 减排", f"{annual_gen_kwh * 0.000015 / 1e4:,.2f} 万吨", "—"],
        ],
        col_widths=[5, 5, 5],
        caption="年度减排指标与 25 年累计",
    )

    _add_heading(doc, "第二章  社会影响", level=1)
    _add_body(doc,
        f"项目建设期间预计创造就业 {int(cap * 5)} 个施工岗位，"
        f"运营期间提供 {int(cap / 10)} 个长期运维岗位。"
        "项目将为当地贡献税收，带动装备制造、运维服务等产业链发展。"
    )

    _add_heading(doc, "第三章  公司治理", level=1)
    _add_styled_table(
        doc, ["治理维度", "实施情况"],
        [
            ["董事会构成", "设立 ESG 委员会，独立董事占比 ≥ 1/3"],
            ["信息披露", "季度 ESG 报告 + 年度可持续发展报告"],
            ["风险管理", "建立三道防线: 业务/风控/审计"],
            ["商业伦理", "反腐败合规培训覆盖率 100%"],
            ["利益相关方", "建立社区沟通机制，定期召开听证会"],
        ],
        col_widths=[5, 10],
        caption="公司治理实施情况",
    )

    _add_heading(doc, "第四章  SDG 对接", level=1)
    _add_styled_table(
        doc, ["SDG 目标", "项目贡献"],
        [
            ["SDG 7 经济适用的清洁能源", f"年提供 {annual_gen_kwh / 1e6:,.0f} GWh 清洁电力"],
            ["SDG 9 产业创新", "推动新能源装备技术升级"],
            ["SDG 13 气候行动", f"年减排 CO₂ {co2_reduction_tons / 1e4:,.1f} 万吨"],
            ["SDG 8 体面工作", f"创造 {int(cap / 10)} 个运维岗位"],
            ["SDG 11 可持续城市", "促进区域能源结构转型"],
        ],
        col_widths=[6, 9],
        caption="SDG 目标对接情况",
    )

    _add_heading(doc, "第五章  持续改进", level=1)
    _add_body(doc,
        "建议建立全生命周期 ESG 管理体系，"
        "定期开展碳排放核查、社会影响评估和治理审计，"
        "持续提升项目 ESG 绩效水平。"
    )


# ── 附录 ──

def _add_appendix(doc, data: Dict, metrics: Dict):
    """附录"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    _add_heading(doc, "附录", level=1)

    _add_heading(doc, "A. 假设条件", level=2)
    fin = data.get("financial", {})
    _add_styled_table(
        doc, ["假设参数", "数值"],
        [
            ["折现率", fmt_pct(fin.get("discount_rate", 0.08))],
            ["电价年涨幅", fmt_pct(fin.get("price_escalation_rate", 0.02))],
            ["运维成本年涨幅", fmt_pct(fin.get("opex_escalation_rate", 0.025))],
            ["组件年衰减率", "0.5%"],
            ["贷款比例", fmt_pct(fin.get("loan_ratio", 0.70))],
            ["贷款利率", fmt_pct(fin.get("loan_interest_rate", 0.045))],
            ["贷款期限", f"{fin.get('loan_term_years', 15)} 年"],
        ],
        col_widths=[6, 10],
        caption="财务模型假设条件",
    )

    _add_heading(doc, "B. 数据来源", level=2)
    sources = data.get("data_sources") or []
    if not sources:
        _add_body(doc, "未附可核验来源；本报告不得用于投资决策。")
    for i, src in enumerate(sources, 1):
        _add_body(doc, _format_citation(i, src))

    _add_heading(doc, "C. 免责声明", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "本报告基于公开资料和项目方提供的数据编制，仅供参考。"
        "报告中的预测和结论受假设条件影响，实际结果可能与预测存在差异。"
        "报告制作方不对依据本报告做出的投资决策承担任何责任。"
        "未经书面许可，不得复制、转发或引用本报告内容。"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_run_fonts(run, FONT_BODY_CN)


# ── 辅助函数 ──

def _project_type_label(t: str) -> str:
    labels = {
        "solar_pv": "光伏发电",
        "onshore_wind": "陆上风电",
        "offshore_wind": "海上风电",
        "storage": "储能",
        "hybrid": "风光储一体化",
        "hydrogen": "绿氢",
    }
    return labels.get(t, t)


def _project_status_label(s: str) -> str:
    labels = {
        "prospecting": "选址阶段",
        "development": "开发阶段",
        "financing": "融资阶段",
        "construction": "建设阶段",
        "commissioning": "调试阶段",
        "operating": "运营阶段",
        "decommissioned": "退役阶段",
    }
    return labels.get(s, s)


def _format_location(data: Dict) -> str:
    parts = []
    if data.get("province_or_region"):
        parts.append(data["province_or_region"])
    if data.get("city") and data["city"] != data.get("province_or_region"):
        parts.append(data["city"])
    return "".join(parts) or "待确认"


# ═══════════════════════════════════════════════════════════════
# 7. PDF 生成器 (reportlab)
# ═══════════════════════════════════════════════════════════════

def generate_pdf(
    data: Dict[str, Any],
    report_type: str,
    template_info: TemplateInfo,
    metrics: Dict[str, Any],
    title: str,
    confidential: str,
    market: str = "cn",
) -> bytes:
    """
    使用 reportlab 生成专业 PDF
    """
    if market in ("global", "en", "int"):
        return _generate_pdf_global(data, report_type, template_info, metrics, title, confidential)
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image as RLImage, KeepTogether,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus.tableofcontents import TableOfContents

    buf = io.BytesIO()

    # 尝试注册中文字体
    cn_font = _register_cn_font()
    cn_font_bold = cn_font  # reportlab 对粗体用同一字体

    primary = colors.HexColor(f"#{COLOR_PRIMARY}")
    accent = colors.HexColor(f"#{COLOR_ACCENT}")
    light_gray = colors.HexColor(f"#{COLOR_LIGHT_GRAY}")
    text_dark = colors.HexColor("#333333")

    # 样式 (字号/间距阶梯取自 LAYOUT, 与 docx 渲染器共用)
    fs = LAYOUT["font_size"]
    sp = LAYOUT["spacing"]
    styles = getSampleStyleSheet()

    style_cover_title = ParagraphStyle(
        "CoverTitle", parent=styles["Title"],
        fontName=cn_font_bold, fontSize=fs["cover_title"], textColor=primary,
        alignment=1, spaceAfter=20,
    )
    style_cover_sub = ParagraphStyle(
        "CoverSub", parent=styles["Normal"],
        fontName=cn_font, fontSize=fs["cover_sub"], textColor=text_dark,
        alignment=1, spaceAfter=10,
    )
    style_h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontName=cn_font_bold, fontSize=fs["h1"], textColor=primary,
        spaceAfter=sp["h1_after"], spaceBefore=sp["h1_before"],
    )
    style_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontName=cn_font_bold, fontSize=fs["h2"], textColor=colors.HexColor("#2C3E50"),
        spaceAfter=sp["h2_after"], spaceBefore=sp["h2_before"],
    )
    style_body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName=cn_font, fontSize=10.5, textColor=text_dark,
        leading=18, spaceAfter=sp["body_after"],
    )
    style_caption = ParagraphStyle(
        "Caption", parent=styles["Normal"],
        fontName=cn_font, fontSize=fs["caption"], textColor=colors.HexColor("#555555"),
        alignment=1, spaceAfter=sp["caption_after"],
    )
    style_small = ParagraphStyle(
        "Small", parent=styles["Normal"],
        fontName=cn_font, fontSize=fs["small"], textColor=colors.HexColor("#888888"),
    )

    story: List[Any] = []

    # ── 封面 ──
    story.append(Spacer(1, 80))
    story.append(Paragraph("━" * 35, ParagraphStyle(
        "Line", fontName=cn_font, fontSize=10, textColor=accent, alignment=1,
    )))
    story.append(Spacer(1, 30))

    type_names = {
        "feasibility": "可行性研究报告",
        "investment": "投资分析报告",
        "compliance": "合规审查报告",
        "esg": "ESG 专项报告",
    }
    story.append(Paragraph(type_names.get(report_type, "项目报告"), style_cover_title))
    story.append(Paragraph(data.get("name", "新能源项目"), style_cover_sub))
    story.append(Spacer(1, 60))

    cover_info = [
        ["项目名称:", data.get("name", "—")],
        ["项目类型:", _project_type_label(data.get("project_type", ""))],
        ["装机容量:", f"{data.get('capacity_mw', 0):.0f} MW"],
        ["项目地点:", _format_location(data)],
        ["报告日期:", datetime.now().strftime("%Y年%m月%d日")],
        ["机密等级:", f"【{_confidential_label_cn(confidential)}】"],
    ]
    cover_table = Table(cover_info, colWidths=[4 * cm, 10 * cm])
    cover_table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), cn_font, 11),
        ("TEXTCOLOR", (0, 0), (0, -1), primary),
        ("TEXTCOLOR", (1, 0), (1, -1), text_dark),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(cover_table)
    story.append(Spacer(1, 60))
    story.append(Paragraph("━" * 35, ParagraphStyle(
        "Line2", fontName=cn_font, fontSize=10, textColor=accent, alignment=1,
    )))
    story.append(PageBreak())

    # ── 目录 ──
    story.append(Paragraph("目  录", style_h1))
    story.append(Spacer(1, 10))
    for i, sec in enumerate(template_info.sections, 1):
        story.append(Paragraph(f"第{i}章  {sec}", style_body))
    story.append(Paragraph("附录", style_small))
    story.append(PageBreak())

    # ── 正文 ──
    pdf_table_counter = [0]  # 「表 N」题注编号, 正文与附录连续
    _build_pdf_content(story, data, metrics, report_type, style_h1, style_h2, style_body, style_caption, cn_font, primary, accent, light_gray, pdf_table_counter)

    # ── 附录 ──
    story.append(PageBreak())
    _build_pdf_appendix(story, data, metrics, style_h1, style_h2, style_body, cn_font, primary, light_gray, style_small, style_caption, pdf_table_counter)

    # ── 页眉页脚: 页眉 = 报告名 + 密级; 页脚 = 页码 ──
    confidential_cn = _confidential_label_cn(confidential)

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFont(cn_font, 8)
        canvas.setFillColor(colors.HexColor("#888888"))
        # 页眉: 左报告名, 右密级
        canvas.drawString(25 * mm, A4[1] - 15 * mm, title)
        canvas.drawRightString(A4[0] - 25 * mm, A4[1] - 15 * mm, f"【{confidential_cn}】")
        # 页脚: 居中页码
        canvas.drawCentredString(A4[0] / 2, 12 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    page = LAYOUT["page"]
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=page["margin_top_cm"] * cm, bottomMargin=page["margin_bottom_cm"] * cm,
        leftMargin=page["margin_left_cm"] * cm, rightMargin=page["margin_right_cm"] * cm,
        title=title, author="新能源智库",
    )
    doc_pdf.build(story, onFirstPage=on_page, onLaterPages=on_page)
    buf.seek(0)
    return buf.read()


def _register_cn_font() -> str:
    """尝试注册中文字体，返回可用字体名"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 按优先级查找系统中文字体
    font_paths = [
        ("/System/Library/Fonts/STHeiti Light.ttc", "STHeiti"),
        ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
        ("/System/Library/Fonts/Hiragino Sans GB.ttc", "HiraginoSansGB"),
        ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WQYZenHei"),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WQYMicroHei"),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
    ]
    for path, name in font_paths:
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            return name
        except Exception:
            continue

    # reportlab 内置 CJK
    try:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        pass

    return "Helvetica"


def _build_pdf_content(story, data, metrics, report_type, style_h1, style_h2, style_body, style_caption, cn_font, primary, accent, light_gray, table_counter):
    """构建 PDF 正文"""
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer

    fin = data.get("financial", {})
    cap = data.get("capacity_mw", 0)

    # 通用表格样式
    def make_table(headers, rows, col_widths=None, caption=None):
        if caption:
            table_counter[0] += 1
            story.append(Paragraph(f"表 {table_counter[0]}  {caption}", style_caption))
        data_rows = [headers] + rows
        t = Table(data_rows, colWidths=[w * cm for w in (col_widths or [8] * len(headers))])
        style_cmds = [
            ("FONT", (0, 0), (-1, -1), cn_font, LAYOUT["font_size"]["table_body"]),
            ("BACKGROUND", (0, 0), (-1, 0), primary),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, 0), LAYOUT["font_size"]["table_head"]),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
        ]
        for i in range(1, len(data_rows)):
            if i % 2 == 0:
                style_cmds.append(("BACKGROUND", (0, i), (-1, i), light_gray))
        t.setStyle(TableStyle(style_cmds))
        return t

    if report_type == "feasibility":
        # Ch1
        story.append(Paragraph("第一章  项目概述", style_h1))
        story.append(Paragraph("1.1 项目背景", style_h2))
        story.append(Paragraph(
            f"本报告对「{data['name']}」进行全面的可行性研究分析。"
            f"项目位于{_format_location(data)}，规划装机容量 {cap:.0f} MW，"
            f"属于{_project_type_label(data['project_type'])}类别。", style_body))
        story.append(Paragraph("1.2 基本信息", style_h2))
        story.append(make_table(
            ["项目参数", "数值"],
            [["项目名称", data["name"]],
             ["装机容量", f"{cap:.0f} MW"],
             ["项目地点", _format_location(data)],
             ["预计寿命", f"{fin.get('project_life_years', 25)} 年"]],
            [6, 10], caption="项目基本信息一览"))
        story.append(Spacer(1, 10))

        # Ch2
        story.append(Paragraph("第二章  资源评估", style_h1))
        annual_gen = fin.get("annual_generation_mwh", 0)
        util_hours = annual_gen / cap if cap > 0 else 0
        story.append(Paragraph(
            f"年等效利用小时数约 {util_hours:,.0f} 小时，首年预计发电量 {annual_gen:,.0f} MWh。",
            style_body))

        # Ch3
        story.append(Paragraph("第三章  技术方案", style_h1))
        story.append(Paragraph(
            f"推荐分块发电、集中并网方案。总装机 {cap:.0f} MW。", style_body))

        # Ch4
        story.append(Paragraph("第四章  财务分析", style_h1))
        story.append(Paragraph("4.1 核心财务指标", style_h2))
        story.append(make_table(
            ["指标", "数值", "评价"],
            [["NPV", fmt_money(metrics["npv"]), "✓" if metrics["npv"] > 0 else "✗"],
             ["IRR", fmt_pct(metrics["irr"]), "✓" if metrics["irr"] > 0.08 else "⚠"],
             ["回收期", f"{metrics['payback_years']:.1f} 年", "✓" if metrics["payback_years"] < 10 else "⚠"],
             ["LCOE", f"{metrics['lcoe']:.4f} 元/kWh", "✓" if metrics["lcoe"] < 0.35 else "✗"]],
            [5, 5, 5], caption="核心财务指标评价"))
        story.append(Spacer(1, 10))
        story.append(Paragraph("4.2 现金流预测 (前5年)", style_h2))
        cf_rows = [[f"第 {y} 年", fmt_money(metrics["cashflows"][y])]
                    for y in range(1, min(6, len(metrics["cashflows"])))]
        story.append(make_table(["年度", "净现金流"], cf_rows, [6, 10],
                                caption="前5年净现金流预测"))

        # Ch5
        story.append(Paragraph("第五章  风险评估", style_h1))
        risk_data = [["政策风险", "中", "补贴政策调整"],
                      ["资源风险", "低", "气象条件波动"],
                      ["技术风险", "低", "设备故障率"],
                      ["资金风险", "中", "利率上升"]]
        story.append(make_table(["风险", "等级", "说明"], risk_data, [4, 3, 8],
                                caption="主要风险识别与等级评定"))

        # Ch6
        story.append(Paragraph("第六章  结论与建议", style_h1))
        npv_ok = metrics["npv"] > 0
        irr_ok = metrics["irr"] > 0.08
        if npv_ok and irr_ok:
            story.append(Paragraph(
                "项目各项财务指标均达到投资要求，建议推进项目开发。", style_body))
        else:
            story.append(Paragraph(
                "部分指标未达预期，建议优化方案后重新评估。", style_body))

    elif report_type == "investment":
        story.append(Paragraph("第一章  执行摘要", style_h1))
        story.append(Paragraph(
            f"项目装机 {cap:.0f} MW，总投资 {fmt_money(metrics['initial_investment'])}。"
            f"IRR {fmt_pct(metrics['irr'])}，NPV {fmt_money(metrics['npv'])}。", style_body))

        story.append(Paragraph("第二章  市场分析", style_h1))
        story.append(Paragraph(
            "2024年全国光伏新增装机超200GW。预计2025-2030年CAGR 15-20%。", style_body))

        story.append(Paragraph("第三章  财务模型", style_h1))
        story.append(make_table(
            ["核心指标", "数值", "基准"],
            [["NPV", fmt_money(metrics["npv"]), "> 0"],
             ["IRR", fmt_pct(metrics["irr"]), "≥ 8%"],
             ["回收期", f"{metrics['payback_years']:.1f} 年", "≤ 10 年"],
             ["LCOE", f"{metrics['lcoe']:.4f}", "≤ 0.35"]],
            [5, 5, 5], caption="财务模型核心指标与行业基准"))

        story.append(Paragraph("第四章  敏感性分析", style_h1))
        base_irr = metrics["irr"]
        tariff = fin.get("electricity_price", 0.32)
        sens_rows = []
        for delta in [-0.10, -0.05, 0, 0.05, 0.10]:
            adj_fin = {**fin, "electricity_price": tariff * (1 + delta)}
            adj_m = calc_financial_metrics(adj_fin)
            sens_rows.append([f"{delta:+.1%}", fmt_pct(adj_m["irr"]),
                              f"{(adj_m['irr'] - base_irr):+.1%}"])
        story.append(make_table(["电价变动", "IRR", "变化"], sens_rows, [5, 5, 5],
                                caption="电价敏感性分析"))

        story.append(Paragraph("第五章  ESG 评估", style_h1))
        co2 = fin.get("annual_generation_mwh", 0) * 1000 * 0.585 / 1e6
        story.append(Paragraph(
            f"年减排 CO₂ 约 {co2:,.2f} 万吨，ESG 评级 A 级。", style_body))

        story.append(Paragraph("第六章  投资建议", style_h1))
        ok = metrics["npv"] > 0 and metrics["irr"] > 0.08
        story.append(Paragraph(
            "推荐投资。" if ok else "暂缓投资，需优化后重新评估。", style_body))

    elif report_type == "compliance":
        story.append(Paragraph("第一章  法规清单", style_h1))
        story.append(make_table(
            ["法规名称", "适用", "状态"],
            [["《可再生能源法》", "全国", "✓ 符合"],
             ["《电力法》", "全国", "✓ 符合"],
             ["《环境保护法》", "全国", "✓ 符合"],
             ["《土地管理法》", "全国", "⚠ 待确认"]],
            [7, 3, 4], caption="适用法规合规性清单"))

        story.append(Paragraph("第二章  许可证状态", style_h1))
        story.append(make_table(
            ["许可证", "主管部门", "状态", "预计时间"],
            [["项目核准", "发改委", "待办理", "T+3月"],
             ["环评批复", "生态环境局", "待办理", "T+6月"],
             ["土地预审", "自然资源局", "待办理", "T+4月"]],
            [4, 4, 3, 3], caption="许可证办理状态"))

        story.append(Paragraph("第三章  环境影响评估", style_h1))
        story.append(Paragraph(
            "施工期影响: 噪声、扬尘。运营期影响: 电磁辐射、生态。", style_body))

        story.append(Paragraph("第四章  土地使用分析", style_h1))
        story.append(Paragraph(
            f"选址 {_format_location(data)}，需确认土地性质并办理审批。", style_body))

        story.append(Paragraph("第五章  电网接入方案", style_h1))
        story.append(Paragraph(
            f"装机 {cap:.0f} MW，建议接入 110kV/220kV 变电站。", style_body))

    elif report_type == "esg":
        annual_gen_kwh = fin.get("annual_generation_mwh", 0) * 1000
        co2 = annual_gen_kwh * 0.585 / 1e6  # 万吨

        story.append(Paragraph("第一章  碳排放与减排", style_h1))
        story.append(make_table(
            ["指标", "年度", "25年累计"],
            [["发电量", f"{annual_gen_kwh / 1e6:,.1f} GWh",
              f"{annual_gen_kwh / 1e6 * 25:,.0f} GWh"],
             ["CO₂减排", f"{co2:,.2f} 万吨", f"{co2 * 25:,.0f} 万吨"]],
            [5, 5, 5], caption="年度减排指标与 25 年累计"))

        story.append(Paragraph("第二章  社会影响", style_h1))
        story.append(Paragraph(
            f"建设期就业 {int(cap * 5)} 人，运营期 {int(cap / 10)} 人。", style_body))

        story.append(Paragraph("第三章  公司治理", style_h1))
        story.append(make_table(
            ["维度", "情况"],
            [["ESG委员会", "已设立"],
             ["信息披露", "季度+年度"],
             ["风险管理", "三道防线"]],
            [5, 10], caption="公司治理实施情况"))

        story.append(Paragraph("第四章  SDG 对接", style_h1))
        story.append(make_table(
            ["SDG", "贡献"],
            [["SDG 7", f"年提供 {annual_gen_kwh / 1e6:,.0f} GWh 清洁电力"],
             ["SDG 13", f"年减排 {co2:,.1f} 万吨 CO₂"]],
            [5, 10], caption="SDG 目标对接情况"))

        story.append(Paragraph("第五章  持续改进", style_h1))
        story.append(Paragraph(
            "建立全生命周期 ESG 管理体系，持续提升绩效。", style_body))


def _build_pdf_appendix(story, data, metrics, style_h1, style_h2, style_body, cn_font, primary, light_gray, style_small, style_caption=None, table_counter=None):
    """PDF 附录"""
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer

    fin = data.get("financial", {})

    story.append(Paragraph("附录", style_h1))

    story.append(Paragraph("A. 假设条件", style_h2))
    if style_caption is not None and table_counter is not None:
        table_counter[0] += 1
        story.append(Paragraph(f"表 {table_counter[0]}  财务模型假设条件", style_caption))
    rows = [
        ["折现率", fmt_pct(fin.get("discount_rate", 0.08))],
        ["电价年涨幅", fmt_pct(fin.get("price_escalation_rate", 0.02))],
        ["运维成本年涨幅", fmt_pct(fin.get("opex_escalation_rate", 0.025))],
        ["组件年衰减率", "0.5%"],
        ["贷款比例", fmt_pct(fin.get("loan_ratio", 0.70))],
    ]
    t = Table([["参数", "数值"]] + rows, colWidths=[6 * cm, 10 * cm])
    t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), cn_font, LAYOUT["font_size"]["table_body"]),
        ("BACKGROUND", (0, 0), (-1, 0), primary),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))

    story.append(Paragraph("B. 数据来源", style_h2))
    sources = data.get("data_sources") or []
    if not sources:
        story.append(Paragraph("未附可核验来源；本报告不得用于投资决策。", style_body))
    for i, source in enumerate(sources, 1):
        story.append(Paragraph(_format_citation(i, source), style_body))

    story.append(Paragraph("C. 免责声明", style_h2))
    story.append(Paragraph(
        "本报告基于公开资料和项目方提供的数据编制，仅供参考。"
        "报告中的预测和结论受假设条件影响。"
        "未经书面许可不得复制或转发。", style_small))


# ═══════════════════════════════════════════════════════════════
# 8. 报告生成主函数
# ═══════════════════════════════════════════════════════════════

# 内存存储 (无 DB 依赖时的 fallback)
_report_store: Dict[str, Dict[str, Any]] = {}


async def _generate_report_task(
    report_id: str,
    data: Dict[str, Any],
    report_type: str,
    output_format: str,
    title: str,
    confidential: str,
):
    """后台任务: 实际生成报告文件"""
    template_info = TEMPLATES[report_type]
    metrics = calc_financial_metrics(data.get("financial", {}))

    try:
        if output_format == "docx":
            if not HAS_DOCX:
                raise ImportError("python-docx 未安装")
            file_bytes = generate_docx(data, report_type, template_info, metrics, title, confidential)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            if not HAS_REPORTLAB:
                raise ImportError("reportlab 未安装")
            file_bytes = generate_pdf(data, report_type, template_info, metrics, title, confidential)
            mime = "application/pdf"
            ext = "pdf"

        # 保存到文件系统
        reports_dir = os.path.join(os.getcwd(), "generated_reports")
        os.makedirs(reports_dir, exist_ok=True)
        filename = f"{report_id}.{ext}"
        filepath = os.path.join(reports_dir, filename)
        with open(filepath, "wb") as f:
            f.write(file_bytes)

        # 更新状态
        _report_store[report_id].update({
            "status": "completed",
            "progress": 100,
            "size": len(file_bytes),
            "file_path": filepath,
            "download_url": f"/api/v1/reports/{report_id}/file",
            "completed_at": datetime.now().isoformat(),
        })

        # 后台任务必须使用自己的数据库会话；请求会话在响应后已关闭。
        async with AsyncSessionLocal() as db:
            try:
                result = await db.execute(select(Report).where(Report.id == report_id))
                report = result.scalar_one_or_none()
                if report:
                    report.status = "completed"
                    report.file_path = filepath
                    report.updated_at = datetime.now()
                    await db.commit()
            except Exception as e:
                await db.rollback()
                logger.exception("更新报告完成状态失败")
                raise

        logger.info(f"报告 {report_id} 生成完成: {filepath}")

    except Exception as e:
        logger.error(f"报告 {report_id} 生成失败: {e}")
        _report_store[report_id].update({
            "status": "failed",
            "progress": 0,
            "error": str(e),
        })
        async with AsyncSessionLocal() as db:
            result = await db.execute(select(Report).where(Report.id == report_id))
            report = result.scalar_one_or_none()
            if report:
                report.status = "failed"
                report.content = str(e)[:1000]
                report.updated_at = datetime.now()
                await db.commit()


# ═══════════════════════════════════════════════════════════════
# 9. FastAPI 路由
# ═══════════════════════════════════════════════════════════════

@router.get("/templates")
async def get_report_templates(market: str = "cn"):
    """
    获取可用报告模板列表 — 支持 market=cn|global 双语
    """
    is_gl = market in ("global", "en", "int")

    # 专业模板 — 根据 market 返回对应语言
    pro_templates = []
    for key, tpl in TEMPLATES.items():
        if is_gl:
            pro_templates.append({
                "id": tpl.id,
                "name": tpl.name_en,
                "description": tpl.description_en,
                "sections": tpl.sections_en,
                "icon": tpl.icon,
                "report_type": key,
            })
        else:
            pro_templates.append({
                "id": tpl.id,
                "name": tpl.name_zh,
                "description": tpl.description_zh,
                "sections": tpl.sections_zh,
                "icon": tpl.icon,
                "report_type": key,
            })

    # 兼容旧模板
    if is_gl:
        legacy = [
            {"id": "t1", "name": "Financial Analysis", "description": "IRR, NPV, LCOE full analysis", "category": "financial"},
            {"id": "t2", "name": "Resource Assessment", "description": "GHI, generation, site scoring", "category": "resource"},
            {"id": "t3", "name": "Investment Memo", "description": "Feasibility & investment recommendation", "category": "comprehensive"},
            {"id": "t4", "name": "Due Diligence", "description": "Technical, financial, legal DD", "category": "due_diligence"},
            {"id": "t5", "name": "Annual Operations", "description": "Generation, revenue, O&M summary", "category": "annual"},
        ]
    else:
        legacy = [
            {"id": "t1", "name": "财务分析报告", "description": "IRR、NPV、LCOE完整分析", "category": "financial"},
            {"id": "t2", "name": "资源评估报告", "description": "GHI、发电量、站址评分", "category": "resource"},
            {"id": "t3", "name": "投资决策建议书", "description": "可行性分析与投资建议", "category": "comprehensive"},
            {"id": "t4", "name": "项目尽职调查", "description": "技术、财务、法务尽调", "category": "due_diligence"},
            {"id": "t5", "name": "年度运营报告", "description": "发电量、收益、运维总结", "category": "annual"},
        ]

    return {
        "data": {
            "professional": pro_templates,
            "legacy": legacy,
        },
        "meta": {"total": len(pro_templates) + len(legacy), "market": market},
    }


@router.post("/generate")
async def generate_report(
    req: GenerateReportRequest,
    background_tasks: BackgroundTasks,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """
    生成专业投资分析报告

    - 支持 4 种模板: feasibility / investment / compliance / esg
    - 支持两种格式: pdf / docx
    - 仅使用当前用户拥有的真实项目数据
    """
    report_id = str(uuid4())
    template_info = TEMPLATES[req.report_type.value]

    # 自动标题
    title = req.title or f"{template_info.name} — {datetime.now().strftime('%Y%m%d')}"

    # 获取数据
    data = await fetch_project_data(db, req.project_id, user_id)
    if req.report_type in (ReportTemplateType.FEASIBILITY, ReportTemplateType.INVESTMENT):
        validate_financial_evidence(data)

    # 统一证据封装（19 项要求第 1 项）：来源/假设/计算轨迹随响应透出
    evidence = _build_report_evidence(title, data)

    # 初始化状态
    _report_store[report_id] = {
        "report_id": report_id,
        "title": title,
        "report_type": req.report_type.value,
        "format": req.format.value,
        "status": "generating",
        "progress": 0,
        "confidential": req.confidential.value,
        "created_at": datetime.now().isoformat(),
        "project_id": req.project_id,
        "user_id": user_id,
    }

    # 如果 DB 可用，创建记录
    if db:
        try:
            report = Report(
                id=report_id,
                user_id=user_id,
                project_id=req.project_id,
                title=title,
                report_type=req.report_type.value,
                language="zh",
                status="generating",
                data_sources={"sources": data.get("data_sources", []), "generated_at": datetime.now().isoformat()},
            )
            db.add(report)
            await db.commit()
        except Exception as e:
            _report_store.pop(report_id, None)
            await db.rollback()
            logger.exception("创建报告记录失败")
            raise HTTPException(status_code=503, detail="报告任务创建失败") from e

    # 后台生成
    background_tasks.add_task(
        _generate_report_task,
        report_id=report_id,
        data=data,
        report_type=req.report_type.value,
        output_format=req.format.value,
        title=title,
        confidential=req.confidential.value,
    )

    return {
        "data": {
            "report_id": report_id,
            "status": "generating",
            "report_type": req.report_type.value,
            "format": req.format.value,
            "download_url": f"/api/v1/reports/{report_id}/file",
            "status_url": f"/api/v1/reports/{report_id}",
            "message": "报告生成中，请轮询 status_url 获取进度",
        },
        "meta": {"template": template_info.name},
        "evidence": evidence,
    }


def _build_report_evidence(title: str, data: Dict[str, Any]) -> Dict[str, Any]:
    """为报告生成响应组装 evidence_envelope。

    sources 复用项目/财务模型已有的来源元数据（标题/URL/作者/版本/发布日期/
    页码/许可/获取日期）；calculation_trace 在财务输入齐备时给出 NPV/IRR/
    LCOE/回收期四个可复核公式条目，输入不全则不编造轨迹。
    """
    raw_sources = data.get("data_sources") or []
    sources = [source_from_dict(s) for s in raw_sources if isinstance(s, dict)]

    financial = data.get("financial") or {}
    assumptions = [
        f"{key} = {value}"
        for key, value in financial.items()
        if key != "yearly_cashflows" and isinstance(value, (int, float, str))
    ]

    trace: List[Dict[str, Any]] = []
    if REQUIRED_FINANCIAL_INPUTS <= set(financial):
        metrics = calc_financial_metrics(financial)
        inputs = {key: financial.get(key) for key in sorted(REQUIRED_FINANCIAL_INPUTS)}
        trace = [
            {"formula": "NPV = Σ CF_t / (1 + r)^t", "inputs": inputs,
             "result": round(metrics["npv"], 2)},
            {"formula": "IRR: NPV(r) = 0（Newton-Raphson 求解）", "inputs": inputs,
             "result": round(metrics["irr"], 6)},
            {"formula": "LCOE = 全生命周期成本现值 / 总发电量", "inputs": inputs,
             "result": round(metrics["lcoe"], 6)},
            {"formula": "静态回收期 = 累计现金流转正年份", "inputs": inputs,
             "result": metrics["payback_years"]},
        ]

    return build_envelope(
        f"报告生成任务已受理：{title}",
        sources,
        assumptions=assumptions,
        calculation_trace=trace,
        limitations=["报告基于用户提供的项目假设与来源生成，仅供辅助分析，不构成投资建议"],
    )


@router.get("/{report_id}")
async def get_report_status(
    report_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """
    获取报告状态 / 下载信息
    """
    # 先查内存
    await _assert_report_owner(report_id, user_id, db)
    if report_id in _report_store:
        info = _report_store[report_id]
        return {"data": info}

    # 查文件系统
    for ext in ["pdf", "docx"]:
        filepath = _report_file(report_id, ext)
        if os.path.exists(filepath):
            return {
                "data": {
                    "report_id": report_id,
                    "status": "completed",
                    "progress": 100,
                    "size": os.path.getsize(filepath),
                    "download_url": f"/api/v1/reports/{report_id}/file",
                    "file_path": filepath,
                }
            }

    raise HTTPException(status_code=404, detail=f"报告 {report_id} 不存在")


@router.get("/{report_id}/file")
async def download_report_file(
    report_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """
    下载已生成的报告文件
    """
    await _assert_report_owner(report_id, user_id, db)
    # 查文件系统
    for ext in ["pdf", "docx"]:
        filepath = _report_file(report_id, ext)
        if os.path.exists(filepath):
            mime = (
                "application/pdf" if ext == "pdf"
                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            with open(filepath, "rb") as f:
                content = f.read()
            filename = _report_store.get(report_id, {}).get("title", report_id)
            return StreamingResponse(
                io.BytesIO(content),
                media_type=mime,
                headers={
                    "Content-Disposition": f'attachment; filename="{filename}.{ext}"',
                },
            )

    raise HTTPException(status_code=404, detail=f"报告文件 {report_id} 不存在")


# ── 兼容旧接口 ──────────────────────────────────────────────────
# 以下保留原有路由接口，确保前端兼容

async def _assert_report_owner(report_id: str, user_id: str, db: AsyncSession) -> Report:
    result = await db.execute(
        select(Report).where(Report.id == report_id, Report.user_id == user_id)
    )
    report = result.scalar_one_or_none()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在或无权访问")
    return report


@router.get("")
async def list_reports_compat(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    user_id: str = Depends(get_current_user_id),
):
    """列出报告 (兼容接口)"""
    items = [
        item for item in _report_store.values()
        if item.get("user_id") == user_id
    ]
    total = len(items)
    start = (page - 1) * page_size
    end = start + page_size
    return success(data={
            "items": items[start:end],
            "total": total,
            "page": page,
            "page_size": page_size,
        })


@router.post("")
async def create_report_compat(
    _: Dict[str, Any],
    user_id: str = Depends(get_current_user_id),
):
    """Legacy create-report entrypoint; use /reports/generate for generation."""
    raise HTTPException(status_code=422, detail="Use /api/v1/reports/generate")


@router.delete("/{report_id}")
async def delete_report(
    report_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """删除报告"""
    await _assert_report_owner(report_id, user_id, db)

    info = _report_store.pop(report_id, {})
    filepath = info.get("file_path")
    if filepath and os.path.exists(filepath):
        os.remove(filepath)

    return {"data": {"message": "报告已删除", "report_id": report_id}}


# ═══════════════════════════════════════════════════════════════
# 自测
# ═══════════════════════════════════════════════════════════════
# 测试命令:
#   python3 -c "from app.routers.reports import router; print(router.routes)"
#
# 同步生成测试 (不启动 FastAPI):
#   python3 -c "
#   from app.routers.reports import generate_pdf, generate_docx, calc_financial_metrics, _demo_data, TEMPLATES
#   data = _demo_data()
#   metrics = calc_financial_metrics(data['financial'])
#   pdf = generate_pdf(data, 'feasibility', TEMPLATES['feasibility'], metrics, '测试报告', '内部')
#   print(f'PDF size: {len(pdf)} bytes')
#   docx = generate_docx(data, 'feasibility', TEMPLATES['feasibility'], metrics, '测试报告', '内部')
#   print(f'DOCX size: {len(docx)} bytes')
#   "
